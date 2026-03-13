"""
policy_doc_builder.py  —  Generic IT Policy / Standards Document Generator
===========================================================================
Requirements:  pip install python-docx
Usage:         See bottom of file for a worked example, or import and use
               the builder functions in your own script.

Quick start
-----------
    from policy_doc_builder import DocBuilder

    doc = DocBuilder("My Policy Title", "IT Security Policy")
    doc.metadata([
        ("Document Reference", "IT-SEC-XXX-001"),
        ("Version",            "1.0"),
        ("Owner",              "IT Security"),
    ])
    doc.section("1.  Purpose", "This policy defines ...")
    doc.bullets(["Applies to all staff", "All managed devices"])
    doc.save("My_Policy.docx")
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────────────────────
# Default theme  (override by passing theme=dict(...) to DocBuilder)
# ─────────────────────────────────────────────────────────────────────────────
DEFAULT_THEME = dict(
    dark       = RGBColor(0x1F, 0x4E, 0x79),   # headings, title
    mid        = RGBColor(0x2E, 0x75, 0xB6),   # sub-headings, borders
    light_hex  = "D6E4F0",                      # shaded cell fill
    white      = RGBColor(0xFF, 0xFF, 0xFF),
    grey       = RGBColor(0x55, 0x55, 0x55),
    ok_fg      = RGBColor(0x1D, 0x6A, 0x2E),
    ok_bg      = "E8F5E9",
    warn_fg    = RGBColor(0x8B, 0x00, 0x00),
    warn_bg    = "FDECEA",
    border_hex = "CCCCCC",
    alt_row    = "F5F9FD",                      # alternating data row
    font       = "Arial",
)


# ─────────────────────────────────────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def _set_cell_borders(cell, colour="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), colour)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _bottom_border(para, colour="2E75B6", size="12"):
    pPr    = para._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _run(para, text, bold=False, italic=False,
         colour=None, size_pt=11, font="Arial"):
    run           = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    if colour:
        run.font.color.rgb = colour
    return run


def _style_cell(cell, bg=None, border_colour="CCCCCC"):
    if bg:
        _set_cell_bg(cell, bg)
    _set_cell_borders(cell, colour=border_colour)
    _set_cell_margins(cell)


# ─────────────────────────────────────────────────────────────────────────────
# DocBuilder  —  main public interface
# ─────────────────────────────────────────────────────────────────────────────

class DocBuilder:
    """
    Fluent builder for professional Word policy/standards documents.

    Parameters
    ----------
    title       : Main document title (displayed large at top)
    subtitle    : Smaller subtitle beneath the title
    theme       : dict — override any keys from DEFAULT_THEME
    page        : "A4" (default) or "Letter"
    margins_in  : margin size in inches (default 1.0)
    """

    def __init__(self, title, subtitle="", theme=None, page="A4", margins_in=1.0):
        self._doc   = Document()
        self._theme = {**DEFAULT_THEME, **(theme or {})}
        self._font  = self._theme["font"]
        self._title = title

        # Page setup
        w, h = (Cm(21), Cm(29.7)) if page == "A4" else (Inches(8.5), Inches(11))
        m    = Inches(margins_in)
        for s in self._doc.sections:
            s.page_width  = w;  s.page_height   = h
            s.left_margin = m;  s.right_margin  = m
            s.top_margin  = m;  s.bottom_margin = m

        self._content_width = w - 2 * m   # usable width in EMUs

        # Title
        t = self._doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t.paragraph_format.space_before = Pt(12)
        t.paragraph_format.space_after  = Pt(4)
        _run(t, title, bold=True, colour=self._theme["dark"],
             size_pt=22, font=self._font)

        if subtitle:
            s = self._doc.add_paragraph()
            s.alignment = WD_ALIGN_PARAGRAPH.CENTER
            s.paragraph_format.space_after = Pt(16)
            _run(s, subtitle, colour=self._theme["grey"],
                 size_pt=13, font=self._font)

    # ── Content primitives ────────────────────────────────────────────────────

    def heading(self, text, level=1):
        """Add a section heading.  level=1 (large) or level=2 (smaller)."""
        p = self._doc.add_paragraph()
        if level == 1:
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(6)
            col = self._theme["dark"]
            sz  = 14
            _bottom_border(p, colour="%02X%02X%02X" % (
                self._theme["mid"][0],
                self._theme["mid"][1],
                self._theme["mid"][2]))
        else:
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
            col = self._theme["mid"]
            sz  = 12
        _run(p, text, bold=True, colour=col, size_pt=sz, font=self._font)
        return self

    def text(self, body, space_pt=4):
        """Add a plain body paragraph."""
        p = self._doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_pt)
        p.paragraph_format.space_after  = Pt(space_pt)
        _run(p, body, size_pt=11, font=self._font)
        return self

    def bullets(self, items):
        """Add a bullet list.  items = list of strings."""
        for item in items:
            p = self._doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            _run(p, item, size_pt=11, font=self._font)
        return self

    def numbered(self, items):
        """Add a numbered list.  items = list of strings."""
        for item in items:
            p = self._doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            _run(p, item, size_pt=11, font=self._font)
        return self

    def spacer(self):
        """Add a blank paragraph for breathing room."""
        self._doc.add_paragraph()
        return self

    def note(self, text, label="NOTE"):
        """Add a shaded callout / note box (single-cell table)."""
        tbl = self._doc.add_table(rows=1, cols=1)
        c   = tbl.rows[0].cells[0]
        _set_cell_bg(c, self._theme["light_hex"])
        _set_cell_borders(c, colour=self._theme["border_hex"])
        _set_cell_margins(c, top=120, bottom=120, left=160, right=160)
        p = c.paragraphs[0]
        _run(p, f"{label}: ", bold=True, colour=self._theme["dark"],
             size_pt=10, font=self._font)
        _run(p, text, size_pt=10, font=self._font)
        self._doc.add_paragraph()
        return self

    # ── Section shortcut ──────────────────────────────────────────────────────

    def section(self, heading_text, body_text=None, level=1):
        """Heading + optional body paragraph in one call."""
        self.heading(heading_text, level=level)
        if body_text:
            self.text(body_text)
        return self

    # ── Tables ────────────────────────────────────────────────────────────────

    def metadata(self, rows, col_widths=(1.8, 5.0)):
        """
        Two-column key/value metadata table.
        rows = list of (label, value) tuples
        """
        t  = self._theme
        tbl = self._doc.add_table(rows=len(rows), cols=2)
        tbl.style = "Table Grid"
        cw = [Inches(col_widths[0]), Inches(col_widths[1])]

        for i, (label, value) in enumerate(rows):
            cells = tbl.rows[i].cells
            _style_cell(cells[0], bg=t["light_hex"])
            cells[0].width = cw[0]
            _run(cells[0].paragraphs[0], label, bold=True,
                 colour=t["dark"], size_pt=10, font=self._font)

            _style_cell(cells[1])
            cells[1].width = cw[1]
            _run(cells[1].paragraphs[0], value, size_pt=10, font=self._font)

        self._doc.add_paragraph()
        return self

    def simple_table(self, headers, rows, col_widths=None,
                     header_bg=None, alt_rows=True):
        """
        Generic table with a coloured header row and optional alternating rows.

        headers    : list of column header strings
        rows       : list of lists (one inner list per row)
        col_widths : list of widths in inches (auto-distributed if omitted)
        header_bg  : hex string for header background (uses theme dark if omitted)
        alt_rows   : zebra-stripe data rows
        """
        t   = self._theme
        n   = len(headers)
        hbg = header_bg or ("%02X%02X%02X" % (
              t["dark"][0], t["dark"][1], t["dark"][2]))

        # Auto-distribute column widths
        if col_widths is None:
            total  = self._content_width / 914400  # EMU → inches
            each   = round(total / n, 2)
            col_widths = [each] * n

        tbl = self._doc.add_table(rows=1 + len(rows), cols=n)
        tbl.style = "Table Grid"

        # Header
        for j, hdr in enumerate(headers):
            c = tbl.rows[0].cells[j]
            _style_cell(c, bg=hbg)
            c.width = Inches(col_widths[j])
            p = c.paragraphs[0]
            _run(p, hdr, bold=True, colour=t["white"],
                 size_pt=10, font=self._font)

        # Data rows
        for i, row_data in enumerate(rows):
            bg = t["alt_row"] if (alt_rows and i % 2 == 0) else None
            for j, cell_text in enumerate(row_data):
                c = tbl.rows[i + 1].cells[j]
                _style_cell(c, bg=bg)
                c.width = Inches(col_widths[j])
                _run(c.paragraphs[0], str(cell_text),
                     size_pt=10, font=self._font)

        self._doc.add_paragraph()
        return self

    def status_table(self, headers, rows, col_widths=None,
                     status_col=1, approved_values=("APPROVED", "YES", "✓")):
        """
        Table where one column is colour-coded green/red based on its value.

        headers        : list of column header strings
        rows           : list of (col0, col1, ..., colN) tuples
        status_col     : index of the column to colour-code (default 1)
        approved_values: values in status_col that should render green
        """
        t   = self._theme
        n   = len(headers)
        hbg = "%02X%02X%02X" % (t["dark"][0], t["dark"][1], t["dark"][2])

        if col_widths is None:
            total = self._content_width / 914400
            each  = round(total / n, 2)
            col_widths = [each] * n

        tbl = self._doc.add_table(rows=1 + len(rows), cols=n)
        tbl.style = "Table Grid"

        # Header row
        for j, hdr in enumerate(headers):
            c = tbl.rows[0].cells[j]
            _style_cell(c, bg=hbg)
            c.width = Inches(col_widths[j])
            _run(c.paragraphs[0], hdr, bold=True, colour=t["white"],
                 size_pt=10, font=self._font)

        # Data rows
        for i, row_data in enumerate(rows):
            alt_bg = t["alt_row"] if i % 2 == 0 else None
            for j, cell_text in enumerate(row_data):
                c  = tbl.rows[i + 1].cells[j]
                txt = str(cell_text)

                if j == status_col:
                    approved = txt.upper() in [v.upper() for v in approved_values]
                    bg  = t["ok_bg"]   if approved else t["warn_bg"]
                    fg  = t["ok_fg"]   if approved else t["warn_fg"]
                    _style_cell(c, bg=bg)
                    c.width = Inches(col_widths[j])
                    p = c.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _run(p, txt, bold=True, colour=fg,
                         size_pt=10, font=self._font)
                else:
                    _style_cell(c, bg=alt_bg)
                    c.width = Inches(col_widths[j])
                    _run(c.paragraphs[0], txt, size_pt=10, font=self._font)

        self._doc.add_paragraph()
        return self

    def two_col_table(self, rows, col_widths=(2.1, 4.7),
                      header=None, shaded_labels=True):
        """
        Convenience wrapper for a two-column label/description table,
        optionally with a header row.

        rows         : list of (label, description) tuples
        header       : (left_header, right_header) or None
        shaded_labels: alternate shading on label column
        """
        t   = self._theme
        hbg = "%02X%02X%02X" % (t["dark"][0], t["dark"][1], t["dark"][2])

        total_rows = len(rows) + (1 if header else 0)
        tbl = self._doc.add_table(rows=total_rows, cols=2)
        tbl.style = "Table Grid"
        cw = [Inches(col_widths[0]), Inches(col_widths[1])]

        row_offset = 0
        if header:
            for j, hdr in enumerate(header):
                c = tbl.rows[0].cells[j]
                _style_cell(c, bg=hbg)
                c.width = cw[j]
                _run(c.paragraphs[0], hdr, bold=True, colour=t["white"],
                     size_pt=10, font=self._font)
            row_offset = 1

        for i, (label, desc) in enumerate(rows):
            bg = t["light_hex"] if (shaded_labels and i % 2 == 0) else None
            cells = tbl.rows[i + row_offset].cells

            _style_cell(cells[0], bg=bg)
            cells[0].width = cw[0]
            _run(cells[0].paragraphs[0], label, bold=True,
                 size_pt=10, font=self._font)

            _style_cell(cells[1])
            cells[1].width = cw[1]
            _run(cells[1].paragraphs[0], desc, size_pt=10, font=self._font)

        self._doc.add_paragraph()
        return self

    # ── Sign-off block ────────────────────────────────────────────────────────

    def signoff(self, lines=None):
        """
        Add a sign-off / approval block at the end of the document.
        lines = list of strings (defaults to a standard approved-by block)
        """
        if lines is None:
            lines = [
                "Approved by: ________________________     Date: _______________",
                "Role: ________________________________     Signature: ___________",
            ]
        self.spacer()
        for i, line in enumerate(lines):
            p = self._doc.add_paragraph()
            if i == 0:
                p.paragraph_format.space_before = Pt(20)
            _run(p, line, colour=self._theme["grey"], size_pt=10, font=self._font)
        return self

    # ── Save ──────────────────────────────────────────────────────────────────

    def save(self, path="document.docx"):
        """Write the document to disk and print the output path."""
        self._doc.save(path)
        print(f"Saved: {path}")
        return path


# ─────────────────────────────────────────────────────────────────────────────
# Example  —  re-creates the Browser Standardisation Standard
# Run:  python policy_doc_builder.py
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":

    doc = DocBuilder(
        title    = "WEB BROWSER STANDARDISATION STANDARD",
        subtitle = "IT Security Policy",
    )

    doc.metadata([
        ("Document Reference", "IT-SEC-BROWSER-001"),
        ("Version",            "1.0"),
        ("Classification",     "Internal Use Only"),
        ("Policy Owner",       "IT Security"),
        ("Review Date",        "Annually or upon significant change"),
    ])

    # 1. Purpose
    doc.section("1.  Purpose",
        "This standard establishes the approved web browsers for use on all "
        "corporate devices and defines the requirements for browser configuration, "
        "security, and exceptions.")

    # 2. Scope
    doc.section("2.  Scope", "This standard applies to:")
    doc.bullets([
        "All employees, contractors, and third parties accessing corporate systems via a web browser",
        "All corporate-owned and managed devices, including desktops, laptops, and VDI",
        "Any system or service delivered via a web browser interface",
    ])

    # 3. Approved Browsers
    doc.section("3.  Approved Browsers",
        "The following browsers are approved. All others require a formal exception (see Section 6).")

    doc.status_table(
        headers    = ["Browser", "Status", "Notes"],
        col_widths = [1.8, 1.3, 3.7],
        rows = [
            ("Microsoft Edge",        "APPROVED",      "Primary approved browser. Preferred for all corporate use."),
            ("Google Chrome",         "APPROVED",      "Secondary approved browser. Permitted for corporate use."),
            ("Mozilla Firefox",       "NOT PERMITTED", "Not approved for standard use. Exception required."),
            ("Apple Safari",          "NOT PERMITTED", "Not approved for standard use. Exception required."),
            ("Opera / Brave / Other", "NOT PERMITTED", "Not approved for standard use. Exception required."),
        ],
        approved_values = ("APPROVED",),
    )

    # 4. Rationale
    doc.section("4.  Rationale for Standardisation",
        "The decision to standardise on Microsoft Edge and Google Chrome is based on the following.")

    doc.section("4.1  Microsoft Account Integration & Single Sign-On", level=2,
        body_text=(
            "Microsoft Edge supports seamless pass-through authentication using "
            "corporate Microsoft accounts, enabling:"
        ))
    doc.bullets([
        "Automatic sign-in to Microsoft 365 services without repeated credential prompts",
        "Conditional Access policy enforcement through Azure Active Directory",
        "Compliance with organisational IAM controls",
        "Reduced credential exposure by minimising password entry events",
    ])

    doc.section("4.2  Security & Patch Management", level=2,
        body_text="Both browsers receive regular updates and are supported by enterprise tooling:")
    doc.bullets([
        "Limiting browser-based vulnerabilities to a managed, patched set of applications",
        "Enabling centrally enforced configuration via Group Policy (GPO) and Intune",
        "Consistent support for corporate certificate authorities and TLS inspection",
    ])

    doc.section("4.3  Compatibility", level=2,
        body_text=(
            "Corporate applications and third-party SaaS platforms are tested against "
            "approved browsers. Non-approved browsers may result in degraded functionality."
        ))

    # 5. Configuration
    doc.section("5.  Configuration Requirements",
        "The following baseline requirements apply to all approved browsers:")
    doc.bullets([
        "Browsers must be managed via Group Policy or Intune device configuration profiles",
        "Automatic updates must remain enabled and must not be disabled by users",
        "Browser sync must be restricted to the corporate account only",
        "Safe Browsing / SmartScreen must be enabled at all times",
        "Password saving must be disabled; use the corporate password manager",
        "Extension installation must be restricted to IT-approved extensions only",
        "Private / InPrivate browsing must not be used to circumvent corporate controls",
    ])

    # 6. Exceptions
    doc.section("6.  Exception Process",
        "A formal exception must be raised through IT Security before installing or using "
        "a non-approved browser. Exceptions must include:")
    doc.bullets([
        "Business justification for the use of the non-approved browser",
        "Identification of the specific system or application requiring it",
        "Risk assessment and proposed mitigating controls",
        "Approval from the relevant business owner and IT Security",
    ])
    doc.text(
        "Approved exceptions will be time-limited and subject to periodic review. "
        "Unauthorised use may be subject to disciplinary action."
    )

    # 7. Roles
    doc.section("7.  Roles & Responsibilities")
    doc.two_col_table(
        header = ("Role", "Responsibility"),
        rows   = [
            ("IT Security",
             "Own and maintain this standard; review exceptions; enforce compliance"),
            ("IT Operations / Desktop",
             "Deploy and configure approved browsers; apply GPO and Intune settings; manage updates"),
            ("All Staff",
             "Use only approved browsers; report issues; seek exception approval before using alternatives"),
        ],
    )

    # 8. Compliance
    doc.section("8.  Compliance & Enforcement",
        "Compliance is mandatory. IT Security and IT Operations will periodically audit devices. "
        "Non-compliant browsers will be subject to removal. Repeated non-compliance may be "
        "escalated in accordance with the organisation's disciplinary procedures.")

    # 9. Review
    doc.section("9.  Document Review",
        "This standard will be reviewed annually or following any significant change to the "
        "organisation's browser landscape, security requirements, or regulatory obligations.")

    doc.signoff()
    doc.save("Browser_Standardisation_Standard.docx")
