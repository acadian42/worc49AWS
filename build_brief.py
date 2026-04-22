#!/usr/bin/env python3
"""
Build the Local LLM Risk & Control Overview as a one-page Word document.

Requirements:
    pip install python-docx

Usage:
    python build_brief.py
    # Writes local_llm_risk_brief.docx in the current directory.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

# ---- Colors ----
NAVY = RGBColor(0x1F, 0x38, 0x64)
HEADER_BG = "2E5597"       # table header fill (hex, no #)
HIGH_BG = "F4CCCC"
MED_BG = "FFF2CC"
LOW_BG = "E2EFDA"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY_MUTED = RGBColor(0x59, 0x59, 0x59)

# ---- Risk content (verified against primary sources, April 2026) ----
RISKS = [
    ("Supply-chain compromise of model files",
     "Pickle / PyTorch (.pth) model files can execute arbitrary code on load. Active exploitation documented on Hugging Face: malicious models deploying remote access trojans (Rapid7 Labs, 2025); ~100 code-execution models identified (JFrog); PickleScan zero-day bypasses disclosed Dec 2025.",
     "HIGH"),
    ("Data sovereignty via foreign hosted AI",
     "If \"local LLM\" is reinterpreted as pointing at foreign hosted services, exposure follows. DeepSeek's app and API have been shown to transmit user data to ByteDance infrastructure (Volcengine) with storage in the PRC (SecurityScorecard STRIKE, NowSecure, 2025).",
     "HIGH"),
    ("Sensitive data exposure",
     "Agentic LLMs need access to internal data to be useful. Without DLP at the model boundary, prompts, context and outputs can leak regulated or commercially sensitive content.",
     "HIGH"),
    ("Prompt injection & agent hijack",
     "Agents that read emails, documents or web pages can be hijacked by adversarial instructions embedded in that content, triggering unintended tool use, data sharing or destructive actions.",
     "HIGH"),
    ("Excessive agent permissions",
     "Agents granted broad filesystem, identity or API scopes can act outside intended bounds. Risk compounds with prompt injection above.",
     "HIGH"),
    ("Shadow AI & sprawl",
     "Distributed local installs produce no central visibility, no audit trail and no ability to respond to incidents or recall a compromised model.",
     "MEDIUM"),
    ("Compliance & licensing",
     "Several leading open-weight models carry commercial-use restrictions and create data-residency exposure under customer contracts and GDPR.",
     "MEDIUM"),
    ("Endpoint capability",
     "Corporate laptop hardware cannot run models large enough to be meaningfully agentic, which pushes users toward personal devices and unsanctioned accounts.",
     "LOW"),
]

RATING_SHADE = {"HIGH": HIGH_BG, "MEDIUM": MED_BG, "LOW": LOW_BG}


# ---- Low-level XML helpers (python-docx doesn't expose these directly) ----
def set_cell_shading(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color="BFBFBF", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def set_cell_margins(cell, top=20, bottom=20, left=110, right=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tc_pr.append(mar)


def set_table_fixed_layout(table):
    """Force fixed (non-auto-fit) layout so cell widths are honoured."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    # Remove existing tblLayout if any, then add fixed
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    # Also set tblW to the sum of column widths so Word doesn't rescale
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)


def set_column_width(table, col_idx, width_dxa):
    """Set a column width on every cell AND update the tblGrid entry."""
    # Update grid
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        cols = grid.findall(qn("w:gridCol"))
        if col_idx < len(cols):
            cols[col_idx].set(qn("w:w"), str(width_dxa))
    # Update each cell in this column
    for row in table.rows:
        tc = row.cells[col_idx]._tc
        tcPr = tc.get_or_add_tcPr()
        # Remove prior tcW
        for existing in tcPr.findall(qn("w:tcW")):
            tcPr.remove(existing)
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(width_dxa))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)


def add_bottom_border(paragraph, color="2E5597"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:color"), color)
    b.set(qn("w:space"), "4")
    p_bdr.append(b)
    p_pr.append(p_bdr)


def write_cell(cell, text, *, bold=False, color=None, size_pt=9, shade=None,
               align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    if shade:
        set_cell_shading(cell, shade)
    set_cell_borders(cell)
    set_cell_margins(cell)


def heading(doc, text, *, size_pt=11, space_before=5, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size_pt)
    r.font.color.rgb = NAVY
    return p


# ---- Build document ----
def build(output_path="local_llm_risk_brief.docx"):
    doc = Document()

    # Page setup: US Letter, 0.75" margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # --- Title ---
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Risk & Control Overview: Locally Installed LLMs & Agentic AI")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY

    # --- Subtitle / classification line (neutral, replaces prior metadata) ---
    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after = Pt(8)
    add_bottom_border(sub)
    sr = sub.add_run("Security Architecture Brief  \u2022  Internal Use Only  \u2022  Version 1.0")
    sr.italic = True
    sr.font.size = Pt(9)
    sr.font.color.rgb = GREY_MUTED

    # --- Context ---
    ctx = doc.add_paragraph()
    ctx.paragraph_format.space_before = Pt(0)
    ctx.paragraph_format.space_after = Pt(4)
    r1 = ctx.add_run("Context. ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = ctx.add_run(
        "A request has been made to enable selected users to access locally installed LLMs for agentic AI use. "
        "This brief summarises the principal security risks and recommends a centralised enablement path in "
        "preference to distributed local installation."
    )
    r2.font.size = Pt(10)

    # --- Principal Risks ---
    heading(doc, "Principal Risks")

    tbl = doc.add_table(rows=1 + len(RISKS), cols=3)
    tbl.autofit = False

    # Column widths in DXA (1 inch = 1440 DXA). Total = 10080 DXA = 7.0".
    # 1.90" risk | 3.85" description | 1.25" rating
    col_widths_dxa = [2736, 5544, 1800]
    set_table_fixed_layout(tbl)
    for idx, w in enumerate(col_widths_dxa):
        set_column_width(tbl, idx, w)

    hdr = tbl.rows[0].cells
    write_cell(hdr[0], "Risk", bold=True, color=WHITE, shade=HEADER_BG)
    write_cell(hdr[1], "Description", bold=True, color=WHITE, shade=HEADER_BG)
    write_cell(hdr[2], "Rating", bold=True, color=WHITE, shade=HEADER_BG)

    for i, (risk, desc, rating) in enumerate(RISKS, start=1):
        row = tbl.rows[i].cells
        write_cell(row[0], risk, bold=True)
        write_cell(row[1], desc)
        write_cell(row[2], rating, bold=True, shade=RATING_SHADE[rating])

    # --- Recommended Approach ---
    heading(doc, "Recommended Approach \u2014 Centralise, Don\u2019t Distribute")

    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(0)
    lead.paragraph_format.space_after = Pt(4)
    lead.add_run(
        "Deliver agentic AI through a single enterprise AI gateway (indicative: "
    ).font.size = Pt(10)
    g = lead.add_run("ai.company.internal")
    g.bold = True
    g.font.size = Pt(10)
    lead.add_run(
        ") acting as a policy-enforcing proxy in front of both managed frontier models and any internally "
        "hosted open-weight models. Required controls:"
    ).font.size = Pt(10)

    controls = [
        ("Approved model registry only",
         " \u2014 no ad-hoc downloads; frontier models via Vertex or Bedrock; open-weight models hosted internally and vetted."),
        ("Identity-based access",
         " \u2014 SSO, per-user quotas, role-scoped model access."),
        ("DLP and content filtering",
         " \u2014 extend the existing AI Guard / Zscaler pattern to inspect prompts, context and outputs."),
        ("Agent permission scoping",
         " \u2014 explicit tool and data allow-lists; no broad filesystem or identity scopes by default."),
        ("Full audit logging",
         " \u2014 every prompt, tool call and output captured for incident response and review."),
        ("Egress restriction & artifact vetting",
         " \u2014 deny outbound internet from hosted model processes; scan pickle/.pth files before load; prefer Safetensors format."),
    ]
    for label, rest in controls:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        rr = p.add_run(label)
        rr.bold = True
        rr.font.size = Pt(10)
        p.add_run(rest).font.size = Pt(10)

    # --- Asks ---
    heading(doc, "Asks")

    asks = [
        ("Clarify intent.",
         " Confirm whether \u201clocal LLM\u201D means on-endpoint installs, or agent tooling already supported."),
        ("Endorse the gateway pattern",
         " as the enablement vehicle in preference to per-user local installs."),
        ("Approve a bounded pilot",
         " (\u226410 users, approved models only) to validate controls before broader rollout."),
    ]
    for label, rest in asks:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        rr = p.add_run(label)
        rr.bold = True
        rr.font.size = Pt(10)
        p.add_run(rest).font.size = Pt(10)

    doc.save(output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    build()
