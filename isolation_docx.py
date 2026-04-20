"""
Application Isolation Guide - DOCX Generator
=============================================
Generates a CompanyLLC-branded Application Isolation Guide in the same
visual format as the Inspection & Threat Protection Standard, i.e.:

  - Arial throughout
  - Orange cover-page wordmark (#E65100) above an orange rule
  - Heading 1 in orange, Heading 2 in dark gray, Heading 3 in medium gray
  - Italic gray running header with orange bottom rule
  - Footer with "CONFIDENTIAL" (red) and page numbers

Requires:  python-docx  (pip install python-docx)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# =====================================================================
# Brand palette
# =====================================================================
ORANGE       = RGBColor(0xE6, 0x51, 0x00)   # primary accent + H1
DARK_GRAY    = RGBColor(0x33, 0x33, 0x33)   # title / H2
MED_GRAY     = RGBColor(0x55, 0x55, 0x55)   # H3
LIGHT_GRAY   = RGBColor(0x99, 0x99, 0x99)   # metadata / header
RED          = RGBColor(0xCC, 0x00, 0x00)   # confidential marker
BLACK        = RGBColor(0x00, 0x00, 0x00)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

FONT = "Arial"


# =====================================================================
# Low-level XML helpers
# =====================================================================
def set_run(run, *, font=FONT, size=None, bold=None, italic=None, color=None):
    """Apply font attributes to a run."""
    run.font.name = font
    # Ensure east-asian/cs fonts also use Arial so Word doesn't fall back
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


# CT_PPrBase canonical child order (OOXML schema)
_PPR_ORDER = [
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr", "widowControl",
    "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs", "suppressAutoHyphens",
    "kinsoku", "wordWrap", "overflowPunct", "topLinePunct", "autoSpaceDE",
    "autoSpaceDN", "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind",
    "contextualSpacing", "mirrorIndents", "suppressOverlap", "jc", "textDirection",
    "textAlignment", "textboxTightWrap", "outlineLvl", "divId", "cnfStyle",
    "rPr", "sectPr", "pPrChange",
]
_PPR_RANK = {name: i for i, name in enumerate(_PPR_ORDER)}


def _local(el):
    """Return the local (unprefixed) tag name of an lxml element."""
    tag = el.tag
    return tag.split("}", 1)[1] if "}" in tag else tag


def _pPr_insert(pPr, new_el):
    """Insert new_el into pPr at the position required by CT_PPrBase order."""
    new_rank = _PPR_RANK.get(_local(new_el), 999)
    for i, existing in enumerate(pPr):
        if _PPR_RANK.get(_local(existing), 999) > new_rank:
            pPr.insert(i, new_el)
            return
    pPr.append(new_el)


def _pPr_get_or_add(pPr, local_name):
    """Return the named child (creating it in canonical position if missing)."""
    existing = pPr.find(qn(f"w:{local_name}"))
    if existing is not None:
        return existing
    el = OxmlElement(f"w:{local_name}")
    _pPr_insert(pPr, el)
    return el


def add_bottom_border(paragraph, color="E65100", size=6):
    """Add a colored bottom border to a paragraph (used for rules)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = _pPr_get_or_add(pPr, "pBdr")
    # Remove any prior bottom border so we don't double up
    for child in list(pBdr):
        if _local(child) == "bottom":
            pBdr.remove(child)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_top_border(paragraph, color="CCCCCC", size=4):
    """Add a colored top border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = _pPr_get_or_add(pPr, "pBdr")
    for child in list(pBdr):
        if _local(child) == "top":
            pBdr.remove(child)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color)
    pBdr.insert(0, top)


def set_paragraph_spacing(paragraph, before=None, after=None, line=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = _pPr_get_or_add(pPr, "spacing")
    if before is not None:
        spacing.set(qn("w:before"), str(before))
    if after is not None:
        spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._element.append(br)


def add_page_number_field(run):
    """Insert a dynamic PAGE field into a run."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, fld_end):
        run._element.append(el)


# =====================================================================
# Style setup
# =====================================================================
def configure_styles(doc):
    # ---- Normal / default ----
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), FONT)

    # ---- Heading 1 (orange) ----
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = ORANGE

    # ---- Heading 2 (dark gray) ----
    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = DARK_GRAY

    # ---- Heading 3 (medium gray) ----
    h3 = doc.styles["Heading 3"]
    h3.font.name = FONT
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = MED_GRAY


# =====================================================================
# Header & footer
# =====================================================================
def build_header_footer(doc, doc_title):
    section = doc.sections[0]

    # ----- Header -----
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    run = hp.add_run(doc_title)
    set_run(run, size=8, italic=True, color=LIGHT_GRAY)
    add_bottom_border(hp, color="E65100", size=4)

    # ----- Footer -----
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    add_top_border(fp, color="CCCCCC", size=4)

    r1 = fp.add_run("CONFIDENTIAL")
    set_run(r1, size=8, color=RED)
    r_sp = fp.add_run("     Page ")
    set_run(r_sp, size=8, color=LIGHT_GRAY)
    r_pg = fp.add_run()
    set_run(r_pg, size=8, color=LIGHT_GRAY)
    add_page_number_field(r_pg)


# =====================================================================
# Cover page
# =====================================================================
def build_cover_page(doc):
    # Top spacer
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=3000)

    # CompanyLLC wordmark
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=200)
    r = p.add_run("CompanyLLC")
    set_run(r, size=26, bold=True, color=ORANGE)  # 26pt = w:sz 52

    # Orange rule
    rule = doc.add_paragraph()
    add_bottom_border(rule, color="E65100", size=6)
    set_paragraph_spacing(rule, after=200)

    # "Document" label
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=60)
    r = p.add_run("Document")
    set_run(r, size=11, color=LIGHT_GRAY)

    # Title
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=400)
    r = p.add_run("Application Isolation Guide")
    set_run(r, size=20, bold=True, color=DARK_GRAY)

    # Spacer before metadata
    spacer2 = doc.add_paragraph()
    set_paragraph_spacing(spacer2, before=2000)

    # Metadata block: label (bold) + value
    def meta_line(label, value):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=120)
        lbl = p.add_run(f"{label}: ")
        set_run(lbl, size=11, bold=True, color=BLACK)
        val = p.add_run(value)
        set_run(val, size=11, color=BLACK)

    meta_line("Document ID", "DOC-IT-SEC-AIG-002")
    meta_line("Version", "2.0 DRAFT")
    meta_line("Date", "April 2026")
    meta_line("Classification", "Confidential")
    meta_line("Owner", "Global Security Office")

    add_page_break(doc)


# =====================================================================
# Content helpers
# =====================================================================
def add_heading(doc, text, level):
    """Add a styled heading. Level: 1 / 2 / 3."""
    p = doc.add_paragraph()
    if level == 1:
        p.style = doc.styles["Heading 1"]
        set_paragraph_spacing(p, before=360, after=200)
        r = p.add_run(text)
        set_run(r, size=18, bold=True, color=ORANGE)
    elif level == 2:
        p.style = doc.styles["Heading 2"]
        set_paragraph_spacing(p, before=240, after=160)
        r = p.add_run(text)
        set_run(r, size=14, bold=True, color=DARK_GRAY)
    else:
        p.style = doc.styles["Heading 3"]
        set_paragraph_spacing(p, before=200, after=120)
        r = p.add_run(text)
        set_run(r, size=12, bold=True, color=MED_GRAY)
    return p


def add_para(doc, text, italic=False):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=120, line=276)
    # Allow simple **bold** inline markup
    _render_inline(p, text, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=80, line=276)
    _render_inline(p, text)
    return p


def _render_inline(paragraph, text, italic=False):
    """Render text with **bold** and *italic* markers inline."""
    import re
    # Split on **...** first
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            set_run(r, size=11, bold=True, italic=italic, color=BLACK)
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            set_run(r, size=11, italic=True, color=BLACK)
        else:
            r = paragraph.add_run(part)
            set_run(r, size=11, italic=italic, color=BLACK)


# =====================================================================
# Document content (sourced from appisolationv2.md, GxO -> CompanyLLC)
# =====================================================================
def build_body(doc):
    # ---- Overview ----
    add_heading(doc, "Overview", level=1)
    add_para(
        doc,
        "This guide details the step-by-step process for isolating legacy "
        "applications at CompanyLLC, ensuring all risks are addressed and "
        "the right teams are engaged. Each section explains the rationale "
        "behind the requirements and highlights the specific teams "
        "responsible for implementation and review."
    )

    # ---- Section 1: Network Segmentation ----
    add_heading(doc, "1. Network Segmentation", level=1)
    add_heading(doc, "Actions", level=2)
    for item in [
        "Place the legacy application in its own network segment.",
        "Apply strict firewall ACLs to restrict traffic to only essential services.",
        "Enforce ingress and egress controls.",
        "Route all ingress/egress traffic through IPS/IDS and enable SSL decryption (unless it breaks the app).",
        "Deploy a Web Application Firewall (WAF) for any internet-facing web services.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Risks Addressed", level=2)
    for item in [
        "Prevents lateral movement if the application is compromised.",
        "Limits exposure to only necessary services.",
        "Protects against common web vulnerabilities.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Teams to Engage", level=2)
    for item in [
        "Network Engineering (VLAN, ACLs, IPS/IDS, SSL decryption, WAF deployment)",
        "Application Owners (to define essential services)",
    ]:
        add_bullet(doc, item)

    # ---- Section 2: Architecture Mapping & Boundary Definition ----
    add_heading(doc, "2. Architecture Mapping & Boundary Definition", level=1)
    add_heading(doc, "Actions", level=2)
    for item in [
        "Produce a comprehensive architecture diagram of the application **before** any isolation work begins. The diagram must show all servers, databases (and their schemas), application tiers, APIs, thick clients, web interfaces, end-user interfaces (RF devices, kiosks, browsers), administrative interfaces, identity sources, and all ports and protocols used between components.",
        "Maintain a component inventory alongside the diagram, capturing operating system, database versions, vendor support status, and data classification for each asset.",
        "Define a formal **approved cross-bubble service inventory** listing the services allowed to ingress or egress the isolation boundary. Default-approved services include: security tooling (EDR/XDR, vulnerability scanners, patch management), authentication (AD/LDAP), time synchronization (NTP), centralized logging (SIEM forwarders), and PAM brokering. Any service outside this inventory requires documented risk review and InfoSec approval before being permitted.",
        "For every **database link or cross-bubble data flow**, document: source and destination, direction (read-only vs. read/write), port, authentication method, data classification, and compensating controls. Bidirectional write-capable links must be explicitly approved and risk-accepted by InfoSec leadership; reduce to read-only or remove entirely where feasible.",
        "For **thick clients or end-user clients** that reach into the isolated environment from user workstations (e.g., desktop clients making direct database or API calls), document the exposed ports and authentication method, and ensure the client endpoints are covered by current host-based security controls (EDR, patch management, disk encryption). Client endpoints remain outside the bubble; only the minimum required service ports are exposed through the boundary.",
        "Document each **legacy constraint exception** where the application cannot meet a guide requirement. Required fields: the specific requirement that cannot be met, the technical reason (e.g., vendor does not support TLS 1.2+, no MFA capability), the compensating controls applied, and an expiration date tied to the application's modernization roadmap. Common legacy constraints include:",
    ]:
        add_bullet(doc, item)
    # Sub-bullets for legacy constraint examples
    for item in [
        "*Unencrypted protocols* (HTTP, clear-text SOAP, unencrypted DB listeners): mitigate via strict network-layer isolation, TLS termination at an upstream proxy or WAF where feasible, and restriction to pre-approved source zones only.",
        "*Structural shared or local service accounts* required for application function (e.g., bootstrap accounts for kiosk or RF device startup): document the account's scope, confirm it has no administrative privileges, lock the session to a non-interactive or menu-locked shell, and require a second layer of individually-attributable authentication (e.g., personal AD credentials) at the application layer.",
        "*Non-MFA end-user flows* on floor devices: mitigate via device-only VLANs, NAC-validated hardware, and short idle timeouts.",
    ]:
        p = doc.add_paragraph(style="List Bullet 2")
        set_paragraph_spacing(p, after=60, line=276)
        _render_inline(p, item)
    add_bullet(
        doc,
        "Treat the architecture diagram, component inventory, cross-bubble service inventory, and legacy constraint register as living documents. Review and republish whenever the application changes materially, or at minimum annually."
    )
    add_heading(doc, "Risks Addressed", level=2)
    for item in [
        "Prevents incomplete isolation caused by undocumented components, data flows, or dependencies.",
        "Prevents undocumented DB links, thick-client connections, or legacy integrations from silently re-opening the isolation boundary.",
        "Ensures legacy limitations are acknowledged, mitigated with compensating controls, and tied to a remediation timeline rather than accepted indefinitely.",
        "Provides the evidence base required for validation, audit, and risk acceptance.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Teams to Engage", level=2)
    for item in [
        "Application Owners (component inventory, data flows, diagram)",
        "Architecture / Enterprise Architecture (diagram review and approval)",
        "Network Engineering (port and protocol validation)",
        "Database Administrators (DB link and schema inventory)",
        "InfoSec Engineering/GRC (boundary approval, exception review, compensating control validation)",
    ]:
        add_bullet(doc, item)

    # ---- Section 3: System Hardening & Maintenance ----
    add_heading(doc, "3. System Hardening & Maintenance", level=1)
    add_heading(doc, "Actions", level=2)
    for item in [
        "Disable unnecessary services and ports.",
        "Patch all EOL/EOS components to the most secure version available.",
        "Apply compensating WAF/IPS controls where patching is not feasible.",
        "Test and install new compatible patches promptly.",
        "Isolate legacy databases from production clusters and corporate data lakes.",
        "Establish a defined monthly maintenance window for the environment.",
        "Define process to ensure any new high or critical patches available are installed during this maintenance window.",
        "Validate control effectiveness by conducting differential vulnerability scanning (Baseline vs Post-Isolation) to identify any remaining security gaps.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Risks Addressed", level=2)
    for item in [
        "Reduces attack surface.",
        "Minimizes vulnerabilities from outdated components.",
        "Prevents unauthorized access to sensitive data.",
        "Ensures timely remediation of critical vulnerabilities.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Teams to Engage", level=2)
    for item in [
        "Server/OS Administrators (hardening, patching, maintenance window)",
        "Database Administrators (database isolation)",
        "Application Owners (testing patches, compensating controls)",
    ]:
        add_bullet(doc, item)

    # ---- Section 4: Identity and Access Management ----
    add_heading(doc, "4. Identity and Access Management", level=1)
    add_heading(doc, "Actions", level=2)
    for item in [
        "Enforce least privilege access.",
        "Prohibit shared credentials, except where a structural shared account exception has been documented and approved under Section 2.",
        "Implement Role-Based Access Control (RBAC).",
        "Multi-Factor Authentication (MFA) is mandatory for all administrative access and remote connectivity. For standard user access to legacy applications, MFA is required unless operational constraints (e.g., high-velocity warehouse environments) necessitate an exception.",
        "Compensating Controls for Non-MFA Access: Any access flows exempted from MFA must be documented and secured via alternative controls (e.g., restricted strictly to trusted IP zones, NAC validated devices, or accessed solely via an MFA gated Jump/Host/VDI).",
        "Align legacy AD domains with current CORP forest policy, or use local authentication (must follow defined password standard guidelines).",
        "Prefer dedicated, isolated identity solution (e.g., local accounts managed via PAM).",
        "Active Directory trusts with forests that do not meet internal security standards are strictly prohibited, permitting only temporary exceptions to facilitate active migrations. Once the migration concludes, the trust must be immediately disabled to prevent long term security exposure. Any legacy assets that cannot be migrated must be fully isolated and denied any persistent trust relationship with the corporate environment.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Risks Addressed", level=2)
    for item in [
        "Prevents unauthorized access.",
        "Reduces risk from credential theft or misuse.",
        "Ensures only necessary personnel have access.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Teams to Engage", level=2)
    for item in [
        "InfoSec Engineering/Identity Management (RBAC, MFA, AD domain configuration)",
        "InfoSec GRC (password policy enforcement)",
        "Application Owners (defining roles and access needs)",
    ]:
        add_bullet(doc, item)

    # ---- Section 5: Site Connectivity ----
    add_heading(doc, "5. Site Connectivity", level=1)
    add_heading(doc, "Actions", level=2)
    for item in [
        "Connect site-level devices (scanners, printers, automation) via isolated network paths.",
        "Prohibit direct internet access unless explicitly approved and monitored.",
        "Use site-specific firewalls or SD-WAN segmentation.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Risks Addressed", level=2)
    for item in [
        "Prevents external threats from reaching isolated environments.",
        "Limits exposure from site devices.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Teams to Engage", level=2)
    for item in [
        "Site IT Teams (device connectivity)",
        "Network Engineering (firewalls, SD-WAN)",
    ]:
        add_bullet(doc, item)

    # ---- Section 6: Access Controls and Remote Access ----
    add_heading(doc, "6. Access Controls and Remote Access", level=1)
    add_heading(doc, "Actions", level=2)
    for item in [
        "Log and monitor all access to legacy systems.",
        "Route all remote administrative access (including SSH and RDP) through a Privileged Access Management (PAM) solution that enforces MFA, brokers the session, rotates credentials, and records the session for review. PAM brokering directly to the target host satisfies the remote-access requirement for protocols the PAM solution natively supports; a separate jump host is not required for these flows.",
        "Provision a dedicated, hardened jump host for administrative workflows that require GUI tooling the PAM solution cannot broker directly (e.g., Oracle TOAD and similar fat-client database or infrastructure management tools). Access to the jump host itself must be PAM-brokered, and the host must be limited to the minimum tooling required, have no connectivity outside the isolated environment, and enforce application allow-listing and EDR.",
        "Harden and monitor all jump hosts and PAM connectors.",
        "Restrict tooling and access on jump hosts to authorized personnel only.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Risks Addressed", level=2)
    for item in [
        "Prevents unauthorized remote access.",
        "Detects and responds to suspicious activity.",
        "Limits risk from compromised jump hosts.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Teams to Engage", level=2)
    for item in [
        "Security Engineering (PAM)",
        "Infrastructure (jump host setup and configuration)",
        "Application Owners (access requirements)",
    ]:
        add_bullet(doc, item)

    # ---- Section 7: Risk Measurement and Documentation ----
    add_heading(doc, "7. Risk Measurement and Documentation", level=1)
    add_heading(doc, "Actions", level=2)
    for item in [
        "Document residual risk after segmentation.",
        "Coordinate with the InfoSec TVM team to compare the targeted baseline scan of CompanyLLC environment pre and post implementation and validate risk mitigation in CompanyLLC environment.",
        "Include architectural review findings in risk documentation.",
        "Conduct ongoing monitoring and biannual risk measurement updates.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Risks Addressed", level=2)
    for item in [
        "Ensures leadership is aware of and accepts remaining risks.",
        "Provides a basis for ongoing risk management.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Teams to Engage", level=2)
    for item in [
        "InfoSec TVM Team (scanning, risk details)",
        "InfoSec Engineering/GRC (risk documentation, review)",
        "Business Leadership (risk acceptance)",
        "Application Owners (supporting documentation)",
    ]:
        add_bullet(doc, item)

    # ---- Section 8: Validation and Review ----
    add_heading(doc, "8. Validation and Review", level=1)
    add_heading(doc, "Actions", level=2)
    add_para(doc, "InfoSec Engineering/GRC must validate:")
    for item in [
        "Network segmentation",
        "Validate that no unapproved lateral movement paths exist",
        "Architecture diagram, component inventory, and cross-bubble service inventory are complete and current",
        "Legacy constraint exceptions are documented with compensating controls and expiration dates",
        "Identity controls",
        "Database isolation",
        "Site connectivity",
        "Access controls",
        "WAF deployment",
        "Risk documentation and sign-off",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Risks Addressed", level=2)
    for item in [
        "Ensures all controls are in place and effective.",
        "Provides formal review and approval.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Teams to Engage", level=2)
    for item in [
        "InfoSec Engineering/GRC (validation, review)",
        "All previously engaged teams (for evidence and support)",
    ]:
        add_bullet(doc, item)

    # ---- Section 9: Risk Acceptance and Exception Handling ----
    add_heading(doc, "9. Risk Acceptance and Exception Handling", level=1)
    add_heading(doc, "Actions", level=2)
    for item in [
        "All isolation instances and exceptions undergo InfoSec risk review.",
        "IT Leader and Business President must accept documented risk.",
        "Any risk above \u201Clow\u201D to the broader CompanyLLC environment requires additional formal approval.",
        "Any deviation from requirements must be documented, reviewed, and justified before formal exception process.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Risks Addressed", level=2)
    for item in [
        "Ensures only acceptable risks are taken.",
        "Provides a formal process for exceptions.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Teams to Engage", level=2)
    for item in [
        "InfoSec Engineering/GRC (risk review, exception process)",
        "IT Leadership (risk acceptance)",
        "Business President (risk acceptance)",
        "Application Owners (exception documentation)",
    ]:
        add_bullet(doc, item)

    # ---- Section 10: Logging, Monitoring, and Incident Response ----
    add_heading(doc, "10. Logging, Monitoring, and Incident Response", level=1)
    add_heading(doc, "Actions", level=2)
    for item in [
        "Mandate that all security-relevant logs (OS, application, database, network firewall, WAF, jump host, etc.) be forwarded to the CompanyLLC SIEM.",
        "Ensure all systems within the isolated environment are synchronized with the CompanyLLC NTP service for accurate time correlation.",
        "Deploy standard host-based monitoring and security agents (e.g. EDR/XDR, FIM) on all applicable systems.",
        "Define critical alerts (e.g. failed admin logins, new processes spawned, unexpected network traffic, new local accounts) in coordination with the SOC.",
        "Application Owners must work with the SOC to create a specific Incident Response (IR) playbook for this application, identifying key contacts and escalation procedures.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Risks Addressed", level=2)
    for item in [
        "Provides visibility to detect active compromises.",
        "Enables forensic investigation and rapid response.",
        "Reduces attacker dwell time.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Teams to Engage", level=2)
    for item in [
        "Security Operations Center (SOC) (alerting, IR playbook)",
        "InfoSec Engineering (SIEM integration, agent deployment)",
        "Infrastructure Team (NTP issues)",
        "Application Owners (SME support for playbook)",
    ]:
        add_bullet(doc, item)

    # ---- Section 11: Data Protection and Resiliency ----
    add_heading(doc, "11. Data Protection and Resiliency", level=1)
    add_heading(doc, "Actions", level=2)
    for item in [
        "Classify all data stored and processed by the application (e.g. Public, Internal, Confidential).",
        "Ensure all data is encrypted at-rest and in-transit using current CompanyLLC standards, especially when crossing the isolation boundary.",
        "Configure data backups to be stored in a secure, operationally and geographically isolated location (not in the same segment or blast radius).",
        "Validate that data restoration and recovery procedures are documented.",
        "The application's Business Continuity (BCDR) plan must be updated to reflect the new isolated architecture and tested at least annually.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Risks Addressed", level=2)
    for item in [
        "Protects sensitive data from unauthorized disclosure (in-flight or at-rest).",
        "Ensures business critical services can be recovered in the event of a catastrophic failure or ransomware attack.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Teams to Engage", level=2)
    for item in [
        "Data Privacy / GRC (data classification)",
        "BCDR / Resiliency Team (backup strategy, BCDR plan)",
        "InfoSec Engineering (encryption standards, key management)",
        "Database Administrators (backup implementation)",
    ]:
        add_bullet(doc, item)

    # ---- Key Points for the App Team ----
    add_heading(doc, "Key Points for the App Team", level=1)
    for item in [
        "Isolation is a comprehensive process\u2014network segmentation is only one part.",
        "Each step mitigates specific risks that could expose CompanyLLC to compromise or compliance failures.",
        "A complete architecture diagram and component inventory must be produced **before** isolation work begins\u2014you cannot isolate what you have not mapped.",
        "Engage the right teams early to ensure all requirements are met and documented.",
        "InfoSec Engineering/GRC must validate and approve isolation requirements have been met.",
        "Exceptions require formal documentation and approval\u2014do not bypass any requirement without following the process.",
        "Maintenance windows and patching are mandatory for ongoing security.",
    ]:
        add_bullet(doc, item)


# =====================================================================
# Main
# =====================================================================
def fix_settings_zoom(doc):
    """Ensure the <w:zoom> element in settings.xml has the required percent attribute."""
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")


def build_document(output_path):
    doc = Document()

    # Page setup: US Letter, 1" margins
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    configure_styles(doc)
    build_header_footer(doc, "CompanyLLC Application Isolation Guide")
    build_cover_page(doc)
    build_body(doc)
    fix_settings_zoom(doc)

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    build_document("CompanyLLC_Application_Isolation_Guide.docx")
