#!/usr/bin/env python3
"""
Network Security Document Generator
====================================
Generates three professional network security documents:
  1. Network Security Strategy
  2. Network Segmentation Standard v2.0
  3. Inspection & Threat Protection Standard

Usage:
    pip install python-docx
    python generate_docs.py

Optional: Change COMPANY_NAME below to customize for your organization.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# CONFIGURATION - Change these to customize
# ============================================================
COMPANY_NAME = "PlaceHolder"
ACCENT_COLOR = "E65100"  # Orange
ACCENT_RGB = RGBColor(0xE6, 0x51, 0x00)
DARK_TEXT = RGBColor(0x33, 0x33, 0x33)
MID_TEXT = RGBColor(0x55, 0x55, 0x55)
LIGHT_TEXT = RGBColor(0x99, 0x99, 0x99)
RED_TEXT = RGBColor(0xCC, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW_BG = "FFF3E0"
HEADER_BG = ACCENT_COLOR
FONT_NAME = "Arial"
OUTPUT_DIR = "."  # Change to desired output directory


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcMar_existing = tcPr.find(qn('w:tcMar'))
    if tcMar_existing is not None:
        tcPr.remove(tcMar_existing)
    tcPr.append(tcMar)


def set_cell_border(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(borders)


def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level, (size, color, before, after) in {
        1: (18, ACCENT_RGB, 18, 10),
        2: (14, DARK_TEXT, 12, 8),
        3: (12, MID_TEXT, 10, 6),
    }.items():
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = FONT_NAME
        heading_style.font.size = Pt(size)
        heading_style.font.bold = True
        heading_style.font.color.rgb = color
        heading_style.paragraph_format.space_before = Pt(before)
        heading_style.paragraph_format.space_after = Pt(after)


def add_para(doc, text, bold=False, size=11, color=None, alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bold_para(doc, label, text, size=11):
    p = doc.add_paragraph()
    run_b = p.add_run(label)
    run_b.font.name = FONT_NAME
    run_b.font.size = Pt(size)
    run_b.font.bold = True
    run_n = p.add_run(text)
    run_n.font.name = FONT_NAME
    run_n.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    for run in p.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_numbered_bold(doc, label, text):
    p = doc.add_paragraph(style='List Number')
    run_b = p.add_run(label)
    run_b.font.name = FONT_NAME
    run_b.font.size = Pt(11)
    run_b.font.bold = True
    run_n = p.add_run(text)
    run_n.font.name = FONT_NAME
    run_n.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, HEADER_BG)
        set_cell_border(cell)
        set_cell_margins(cell)

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            set_cell_border(cell)
            set_cell_margins(cell)
            if row_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    doc.add_paragraph()
    return table


def add_severity_table(doc, headers, rows, col_widths=None):
    severity_colors = {
        "Critical": "FFCDD2",
        "High": "FFF9C4",
        "Medium": "C8E6C9",
        "Low": "BBDEFB",
    }
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, HEADER_BG)
        set_cell_border(cell)
        set_cell_margins(cell)

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            if col_idx == 0:
                run.font.bold = True
            set_cell_border(cell)
            set_cell_margins(cell)
            if col_idx == 0 and cell_text in severity_colors:
                set_cell_shading(cell, severity_colors[cell_text])

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    doc.add_paragraph()
    return table


def add_title_page(doc, title, subtitle=None, doc_id="", version="", date="April 2026",
                   classification="Confidential", owner="Global Security Office",
                   supersedes=None):
    for _ in range(6):
        doc.add_paragraph()

    add_para(doc, COMPANY_NAME, bold=True, size=26, color=ACCENT_RGB)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="{ACCENT_COLOR}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    if subtitle:
        add_para(doc, "Document", size=11, color=LIGHT_TEXT, space_after=3)
    add_para(doc, title, bold=True, size=22, color=DARK_TEXT, space_after=4)
    if subtitle:
        add_para(doc, subtitle, size=18, color=RGBColor(0x66, 0x66, 0x66), space_after=20)

    add_para(doc, "CONFIDENTIAL", bold=True, size=12, color=RED_TEXT, space_after=6)

    for _ in range(8):
        doc.add_paragraph()

    add_bold_para(doc, "Document ID: ", doc_id)
    add_bold_para(doc, "Version: ", version)
    add_bold_para(doc, "Date: ", date)
    add_bold_para(doc, "Classification: ", classification)
    if supersedes:
        add_bold_para(doc, "Supersedes: ", supersedes)
    add_bold_para(doc, "Owner: ", owner)

    doc.add_page_break()


def setup_header_footer(doc, header_text):
    for section in doc.sections:
        section.different_first_page_header_footer = True

        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        run = hp.add_run(header_text)
        run.font.name = FONT_NAME
        run.font.size = Pt(8)
        run.font.color.rgb = LIGHT_TEXT
        run.font.italic = True
        pPr = hp._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="{ACCENT_COLOR}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ""
        run_conf = fp.add_run("CONFIDENTIAL")
        run_conf.font.name = FONT_NAME
        run_conf.font.size = Pt(8)
        run_conf.font.color.rgb = RED_TEXT
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = fp._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="8" w:space="1" w:color="CCCCCC"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)


def new_doc():
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    setup_styles(doc)
    return doc


C = COMPANY_NAME


# ============================================================
# DOCUMENT 1: NETWORK SECURITY STRATEGY
# ============================================================

def create_strategy():
    doc = new_doc()

    add_title_page(
        doc,
        "Network Security & Segmentation",
        subtitle="Strategy & Implementation Plan",
        doc_id="DOC-IT-IS-NSS-001",
        version="1.0 DRAFT",
        owner="Global Security Office / Information Security",
    )

    # -- Executive Summary --
    doc.add_heading("Executive Summary", level=1)
    add_para(doc, (
        f"This document establishes the {C} Network Security & Segmentation Strategy, "
        f"providing an end-to-end framework for where and how security controls are enforced "
        f"across the {C} network. It addresses the critical alignment gap identified between "
        f"Network Engineering, Cloud Architecture, and Information Security."
    ))
    add_para(doc, (
        "The strategy introduces a tiered enforcement model that eliminates redundant inspection, "
        "optimizes application performance, and meets security requirements. Each enforcement "
        "point has a clearly defined role so that no traffic flow is inspected at the same depth "
        "in more than one place."
    ))

    doc.add_heading("Problem Statement", level=2)
    add_para(doc, f"{C} currently faces several challenges with its network segmentation approach:")
    add_bullet(doc, (
        "The existing Network Segmentation Standard is outdated and does not reflect the "
        "current hybrid cloud architecture (GCVE, GCE, Equinix hubs, SD-WAN sites)."
    ))
    add_bullet(doc, (
        "No end-to-end policy exists that defines which security controls should be enforced "
        "at which enforcement points, resulting in duplication of firewall policies across "
        "site, hub, and cloud layers."
    ))
    add_bullet(doc, (
        "Duplicated deep packet inspection at multiple transit points creates unnecessary "
        "latency impacting warehouse management and latency-sensitive applications."
    ))
    add_bullet(doc, (
        "Teams lack clarity on which controls exist where, making troubleshooting complex and slow."
    ))
    add_bullet(doc, (
        "The Site Segmentation Standard covers warehouse/office segmentation but does not "
        "integrate with hub or cloud control planes."
    ))

    doc.add_heading("Strategic Objectives", level=2)
    add_numbered(doc, "Define a single, authoritative enforcement model from site to cloud with clear ownership at each tier.")
    add_numbered(doc, "Eliminate redundant inspection to reduce latency and simplify troubleshooting.")
    add_numbered(doc, "Maintain or improve security posture while migrating workloads from data centers to Google Cloud.")
    add_numbered(doc, "Provide a scalable, repeatable model for rapid warehouse deployment.")
    add_numbered(doc, "Write standards that define security requirements independently of enforcement technology, so they remain valid as the infrastructure evolves.")

    doc.add_page_break()

    # -- Tiered Enforcement Model --
    doc.add_heading("Tiered Enforcement Model", level=1)
    add_para(doc, (
        "The core of this strategy is a four-tier enforcement model. Each tier has a defined "
        "role, and security controls are placed at the tier where they deliver the most value "
        "with the least performance impact. The guiding principle is: inspect once at the "
        "optimal point, not everywhere along the path."
    ))

    doc.add_heading("Tier Overview", level=2)
    add_table(doc,
        ["Tier", "Location", "Primary Security Function", "Inspection Level"],
        [
            ["Tier 1", "Warehouse / Office Site",
             "Zone separation (VLAN-based), stateful ACLs for north/south traffic control, east/west controls for automation-specific segments",
             "Stateful Packet Inspection (SPI)"],
            ["Tier 2", "Hub / Equinix Colo",
             "Deep packet inspection, IPS/IDS, anti-malware scanning, service-based application rules, threat intelligence feeds",
             "Deep Packet Inspection (DPI) with IPS"],
            ["Tier 3", "Cloud Edge (GCVE / GCE)",
             "Micro-segmentation at the resource level, application-aware rules closest to the workload",
             "Application-layer micro-segmentation"],
            ["Tier 0", "Internet Edge (Cloud Proxy)",
             "SWG for internet-bound and SaaS traffic; URL filtering, SSL inspection, malware scanning for web traffic",
             "Cloud-proxy DPI for web/SaaS traffic"],
        ],
        col_widths=[0.8, 1.3, 2.8, 1.6],
    )

    doc.add_heading("Tier 1: Site-Level Controls", level=2)
    add_para(doc, (
        f"Warehouse and office sites represent the outermost perimeter of the {C} network. "
        f"Given the scale of {C}\u2019s site portfolio (hundreds of warehouses globally), "
        "site-level controls must be standardized, repeatable, and lightweight to avoid "
        "impacting warehouse operations."
    ))

    doc.add_heading("Enforcement Approach", level=3)
    add_bullet(doc, (
        "Stateful packet inspection on the site firewall (SD-WAN appliance or MPLS perimeter "
        "router with ACL capability). This is cost-effective equipment that does not require "
        "advanced threat protection licensing."
    ))
    add_bullet(doc, (
        "Standard north/south ACL applied uniformly: permit DC and cloud ranges on all ports, "
        "deny RFC 1918 (block inter-site), permit internet via cloud proxy."
    ))
    add_bullet(doc, (
        f"VLAN-based zone separation per the Site Segmentation Standard: {C} Corporate, "
        f"{C} Operations, 3rd Party, External DMZ (where required), and Internet Only (guest)."
    ))
    add_bullet(doc, (
        "East/west controls applied only where site-specific automation or IIoT segmentation "
        "is required (bespoke per-site configuration)."
    ))
    add_bullet(doc, (
        "No IPS or DPI at the site level. This reduces latency, simplifies troubleshooting, "
        "and avoids the need for advanced threat protection licenses on every site appliance."
    ))
    add_bold_para(doc, "Rationale: ", (
        "The site firewall\u2019s role is access control and zone enforcement, not threat inspection. "
        "All traffic leaving the site traverses the hub where DPI/IPS occurs, so duplicating "
        "this at the site adds latency without additional security value. Stateful inspection "
        "is sufficient to enforce zone boundaries and prevent unauthorized lateral movement."
    ))

    doc.add_heading("Tier 2: Hub-Level Controls (Equinix Colocation)", level=2)
    add_para(doc, (
        "The hub layer at Equinix is the primary inspection point for all traffic transiting "
        "between sites and the cloud/data center. This is where the most capable firewall "
        "hardware resides and where deep inspection delivers maximum value."
    ))

    doc.add_heading("Enforcement Approach", level=3)
    add_bullet(doc, "Deep Packet Inspection (DPI) on all traffic passing through the hub. This is the single point of full inspection on the north/south path.")
    add_bullet(doc, "IPS/IDS enabled with regularly updated threat signatures and threat intelligence feeds.")
    add_bullet(doc, "Anti-malware scanning for file transfers traversing the hub.")
    add_bullet(doc, "Service-based application rules governing which applications are accessible from which site classifications.")
    add_bullet(doc, (
        "As workloads migrate from data centers to GCVE/GCE, the service-based rules at the hub "
        "should be simplified to permit-any for cloud-bound traffic (since micro-segmentation "
        "handles granularity in the cloud), but IPS and anti-malware scanning are retained."
    ))
    add_bold_para(doc, "Rationale: ", (
        "The hub has the most capable hardware, is sized for inspection throughput, and sits at "
        "the natural choke point between all sites and all cloud/DC environments. Concentrating "
        "DPI/IPS here means a single point of inspection, a single team responsible for IPS "
        "tuning, and a single place to troubleshoot threat detection issues."
    ))

    doc.add_heading("Tier 3: Cloud-Level Controls (GCVE & GCE)", level=2)
    add_para(doc, (
        "The cloud tier applies the most granular controls, closest to the workloads themselves. "
        "Micro-segmentation provides defense-in-depth without duplicating the DPI/IPS already "
        "performed at the hub."
    ))

    doc.add_heading("GCVE (VMware Cloud on Google)", level=3)
    add_bullet(doc, "NSX distributed firewall provides micro-segmentation at the individual VM level.")
    add_bullet(doc, "Rules are application-aware and enforce server-to-server communication policies (e.g., web tier to app tier to DB tier).")
    add_bullet(doc, "No additional IPS/DPI at the NSX layer unless a specific compliance requirement demands it. The hub already inspects inbound traffic.")

    doc.add_heading("GCE (Native Google Compute)", level=3)
    add_bullet(doc, "Google Cloud VPC firewall rules and/or Cloud Armor applied at the resource level.")
    add_bullet(doc, "Security groups/tags used to enforce least-privilege communication between services.")
    add_bullet(doc, (
        "When GCVE workloads migrate to native GCE, the equivalent micro-segmentation rules "
        "must be translated from NSX to VPC firewall rules as part of the migration runbook."
    ))
    add_bold_para(doc, "Key Note \u2013 GCVE to GCE Migration: ", (
        "When applications move from GCVE (NSX micro-segmentation) to native GCE, there is a "
        "change in the enforcement mechanism but not in the security policy. The application\u2019s "
        "communication matrix must be translated into VPC firewall rules or equivalent. This "
        "must be a formal step in every migration plan."
    ))

    doc.add_heading("Tier 0: Internet Edge (Cloud Proxy)", level=2)
    add_para(doc, (
        "All internet-bound traffic from sites is routed through the cloud proxy (e.g., Zscaler), "
        "providing Secure Web Gateway (SWG) capabilities. This tier handles a distinct traffic "
        "class and operates outside the traditional network perimeter."
    ))
    add_bullet(doc, "The cloud proxy provides URL filtering, cloud DLP, malware scanning, and SSL inspection for internet-bound traffic.")
    add_bullet(doc, "Internet-bound traffic does NOT need to be inspected at the hub for IPS/DPI purposes\u2014it is handled by the cloud proxy.")
    add_bullet(doc, "Internet egress traffic does not traverse the hub at all, reducing hub load and eliminating unnecessary inspection.")

    doc.add_page_break()

    # -- Enforcement Point Matrix --
    doc.add_heading("Enforcement Point Matrix", level=1)
    add_para(doc, "The following matrix defines exactly which security controls are applied at each enforcement point. This is the authoritative reference for avoiding duplication and for troubleshooting.")

    add_table(doc,
        ["Security Control", "Site FW", "Hub FW (Equinix)", "Cloud (NSX/VPC)", "Cloud Proxy", "TSG DC"],
        [
            ["Stateful Packet Inspection", "YES \u2013 Primary", "YES", "N/A", "N/A", "YES"],
            ["Deep Packet Inspection", "NO", "YES \u2013 Primary", "NO", "YES (web)", "YES"],
            ["IPS/IDS", "NO", "YES \u2013 Primary", "Optional", "YES (web)", "YES"],
            ["Anti-Malware", "NO", "YES \u2013 Primary", "NO", "YES (web)", "YES"],
            ["Micro-Segmentation", "NO", "NO", "YES \u2013 Primary", "NO", "N/A"],
            ["Zone Separation (VLAN)", "YES \u2013 Primary", "N/A", "N/A", "N/A", "YES"],
            ["Service-Based App Rules", "NO", "Transitional", "YES \u2013 Primary", "NO", "Transitional"],
            ["URL/Web Filtering", "NO", "NO", "NO", "YES \u2013 Primary", "NO"],
            ["SSL Inspection", "NO", "Optional", "NO", "YES \u2013 Primary", "NO"],
            ["DNS Security", "NO", "YES", "NO", "YES", "YES"],
        ],
        col_widths=[1.6, 0.9, 1.1, 1.0, 0.9, 0.8],
    )
    add_para(doc, (
        "\u201cPrimary\u201d indicates the enforcement point of record. \u201cTransitional\u201d indicates controls "
        "that will be deprecated as cloud migration completes. \u201cOptional\u201d indicates controls "
        "that may be enabled for specific compliance scenarios but are not mandated by default."
    ))

    doc.add_page_break()

    # -- Traffic Flow Walkthrough --
    doc.add_heading("Traffic Flow Walkthrough", level=1)
    add_para(doc, "This section walks through common traffic flows and identifies exactly which checkpoints are traversed and what inspection occurs at each.")

    doc.add_heading("Flow 1: Warehouse User Accessing WMS in GCVE", level=2)
    add_numbered_bold(doc, "Site firewall (Tier 1): ", "Stateful ACL permits traffic to DC/cloud ranges. No DPI. Latency impact: negligible (<1ms).")
    add_numbered_bold(doc, "Hub firewall at Equinix (Tier 2): ", "DPI + IPS inspection. Anti-malware scan. Single deep inspection point. Latency impact: 5\u201315ms depending on traffic volume and rule complexity.")
    add_numbered_bold(doc, "GCVE NSX (Tier 3): ", "Micro-segmentation permits this user\u2019s source to access the WMS application on the required ports. No DPI re-inspection. Latency impact: negligible.")
    add_bold_para(doc, "Total inspection points: ", "3 (stateful, deep, micro-seg). No duplication of DPI.")

    doc.add_heading("Flow 2: Warehouse User Accessing the Internet", level=2)
    add_numbered_bold(doc, "Site firewall (Tier 1): ", "Stateful ACL permits internet access via cloud proxy. Traffic is tunneled directly to the proxy, not via the hub.")
    add_numbered_bold(doc, "Cloud Proxy (Tier 0): ", "Full DPI, URL filtering, SSL inspection, malware scanning. This traffic does NOT traverse the hub.")
    add_bold_para(doc, "Total inspection points: ", "2 (stateful, cloud proxy). Hub is not involved.")

    doc.add_heading("Flow 3: Site-to-Site Communication (Exception-Based)", level=2)
    add_para(doc, "Per existing policy, site-to-site communication is denied by default (RFC 1918 deny rule). Where a business exception exists:")
    add_numbered_bold(doc, "Source site firewall: ", "Stateful ACL with specific exception rule permitting traffic to the destination site\u2019s service.")
    add_numbered_bold(doc, "Hub firewall: ", "DPI + IPS inspection of the cross-site traffic.")
    add_numbered_bold(doc, "Destination site firewall: ", "Stateful ACL permits inbound from the approved source.")
    add_bold_para(doc, "Goal: ", "Eliminate these exceptions over time. Where cross-site access is needed, the application should be accessed via the cloud rather than direct site-to-site, removing the east-west spread risk.")

    doc.add_page_break()

    # -- Document Deliverables --
    doc.add_heading("Document Deliverables", level=1)
    add_para(doc, "This strategy is supported by the following standards and procedures:")
    add_table(doc,
        ["Document", "Description", "Status", "Owner"],
        [
            ["Network Segmentation Standard v2.0",
             "Replaces the legacy standard. Defines network zones, security levels, inter-zone rules, and the tiered enforcement model.",
             "DRAFT", "InfoSec / GSO"],
            ["Site Segmentation Standard v2.0",
             "Updates v1.3. Aligns warehouse/site zones with the tiered model. Adds clarity on Advanced/Premium implementation tiers.",
             "DRAFT", "InfoSec / GSO"],
            ["Inspection & Threat Protection Standard",
             "New document. Defines IPS/IDS, DPI, anti-malware requirements, inspection profiles, and placement rules.",
             "DRAFT", "InfoSec / GSO"],
            ["Cloud Segmentation Standard",
             "New document. Defines micro-segmentation requirements for GCVE (NSX) and GCE (VPC), including migration rules.",
             "Planned", "InfoSec / Cloud Arch"],
        ],
        col_widths=[1.8, 2.6, 0.8, 1.0],
    )

    doc.add_page_break()

    # -- Implementation Plan --
    doc.add_heading("Implementation Plan", level=1)

    doc.add_heading("Phase 1: Foundation (Weeks 1\u20134)", level=2)
    add_bullet(doc, "Finalize and publish the Network Segmentation Standard v2.0 and Site Segmentation Standard v2.0.")
    add_bullet(doc, "Circulate the Enforcement Point Matrix to all stakeholders (Network Engineering, Cloud Architecture, InfoSec, managed service providers).")
    add_bullet(doc, "Gain formal sign-off from the GSO on the tiered enforcement model.")
    add_bullet(doc, "Begin drafting the Inspection & Threat Protection Standard.")

    doc.add_heading("Phase 2: Hub Alignment (Weeks 5\u201312)", level=2)
    add_bullet(doc, "Audit current hub firewall rules against the enforcement matrix. Identify and remediate any site-level rules that duplicate hub-level controls.")
    add_bullet(doc, "Standardize IPS profiles across both Equinix hub locations.")
    add_bullet(doc, "Validate that service-based rules at hubs are correctly transitioning to cloud micro-segmentation as workloads migrate.")
    add_bullet(doc, "Coordinate with managed service providers on aligning NSX micro-segmentation policies with the published standards.")

    doc.add_heading("Phase 3: Site Rollout (Weeks 8\u201324)", level=2)
    add_bullet(doc, "Deploy standard ACLs to all sites per the Site Segmentation Standard.")
    add_bullet(doc, "Upgrade sites designated as Advanced or Premium to include east/west intra-site controls.")
    add_bullet(doc, "Validate that all sites are correctly routing through the hub for DPI/IPS and through the cloud proxy for internet traffic.")

    doc.add_heading("Phase 4: Cloud Alignment (Weeks 12\u201324)", level=2)
    add_bullet(doc, "Publish the Cloud Segmentation Standard covering GCVE and GCE.")
    add_bullet(doc, "Validate that NSX micro-segmentation rules align with the published standard for all migrated workloads.")
    add_bullet(doc, "Establish the GCVE-to-GCE migration security review process to ensure micro-segmentation rules are translated correctly during native cloud migrations.")
    add_bullet(doc, "Audit and set expiry dates on all legacy site-to-site exceptions.")

    doc.add_page_break()

    # -- Risks --
    doc.add_heading("Risks and Considerations", level=1)
    add_table(doc,
        ["Risk", "Description", "Mitigation"],
        [
            ["Single point of inspection failure",
             "Concentrating DPI/IPS at the hub means hub outage removes threat inspection.",
             "Hub redundancy across East and Central sites. Failover ensures continuous inspection."],
            ["Network team capacity",
             "The network team may lack bandwidth to implement all phases in parallel with site rollouts.",
             "Phases are sequenced to avoid overloading. Hub alignment and site rollout overlap only where teams are independent."],
            ["GCVE to GCE migration gap",
             "Micro-segmentation rules in NSX may not be properly translated when workloads move to native GCE.",
             "Require a security review step in every GCVE-to-GCE migration runbook."],
            ["Stakeholder misalignment",
             "Different teams may interpret the tiered model differently.",
             "The Enforcement Point Matrix is the single source of truth. Any deviation requires a security exception."],
            ["Legacy site-to-site exceptions",
             "Existing cross-site exceptions create east-west risk that the new model aims to eliminate.",
             "Audit all exceptions. Set expiry dates. Migrate to cloud-based access patterns."],
            ["Hub capacity growth",
             "Concentrating all DPI/IPS at the hub requires adequate sizing as sites and bandwidth increase.",
             "Hub firewall utilization monitored continuously. Capacity planning triggered at 70% sustained utilization."],
        ],
        col_widths=[1.6, 2.6, 2.3],
    )

    # -- Future Considerations --
    doc.add_heading("Future Considerations", level=1)
    add_para(doc, (
        "The security industry is broadly moving toward Security Service Edge (SSE) and Secure "
        "Access Service Edge (SASE) architectures, which converge networking and security into "
        f"cloud-delivered services. {C} already uses a cloud proxy for Secure Web Gateway "
        "(SWG) capabilities, which is a foundational component of an SSE stack."
    ))
    add_para(doc, (
        "A full SSE/SASE evaluation is not in scope for this strategy. The immediate priority "
        "is closing the alignment gap and establishing the tiered enforcement model using current "
        "infrastructure. However, this strategy has been deliberately written to define security "
        "requirements (what must be inspected) independently of enforcement technology (which "
        "device does the inspecting). This means the standards published under this strategy "
        "will remain valid if the organization later shifts inspection from hub-based firewalls "
        "to cloud-delivered security services."
    ))
    add_para(doc, (
        "When the tiered model is operational and the cloud migration has stabilized, a separate "
        "evaluation of SSE/SASE capabilities (particularly ZTNA as a VPN replacement and CASB "
        "for SaaS protection) may be warranted. That evaluation should be scoped as its own "
        "initiative with its own business case."
    ))

    doc.add_page_break()

    # -- Approval --
    doc.add_heading("Review and Approval", level=1)
    add_table(doc,
        ["Role", "Name", "Date", "Signature"],
        [
            ["CISO / GSO Lead", "", "", ""],
            ["VP Network Engineering", "", "", ""],
            ["Director Cloud Architecture", "", "", ""],
            ["InfoSec Standards Lead", "", "", ""],
        ],
        col_widths=[2.0, 1.8, 1.4, 1.4],
    )

    setup_header_footer(doc, f"{C} Network Security Strategy")
    return doc


# ============================================================
# DOCUMENT 2: NETWORK SEGMENTATION STANDARD v2.0
# ============================================================

def create_netseg_standard():
    doc = new_doc()

    add_title_page(
        doc,
        "Network Segmentation Standard",
        doc_id="DOC-IT-IS-NSS-002",
        version="2.0 DRAFT",
        supersedes="Network Segmentation Standard (legacy, undated)",
        owner="Global Security Office",
    )

    # 1. Executive Summary
    doc.add_heading("Executive Summary", level=1)
    add_para(doc, (
        "Network Segmentation is the practice of dividing a network into isolated zones and "
        "controlling communications between them using policy-based rules. Effective segmentation "
        "limits the blast radius of a security incident, makes lateral movement more difficult "
        "for adversaries, and increases the likelihood of detecting unauthorized activity."
    ))
    add_para(doc, (
        f"This standard replaces the legacy {C} Network Segmentation Standard and provides a "
        "comprehensive framework covering site networks, hub/colocation facilities, data centers, "
        f"and cloud environments. It is designed to work in conjunction with the {C} Site "
        f"Segmentation Standard and the {C} Network Security Strategy."
    ))

    # 2. Purpose
    doc.add_heading("Purpose", level=1)
    add_para(doc, (
        f"This standard defines the network zones used for segmentation across the {C} network, "
        "the criteria for system placement within these zones, the rules governing communications "
        "between zones, and the security levels assigned to each zone."
    ))

    # 3. Scope
    doc.add_heading("Scope", level=1)
    add_para(doc, (
        f"This standard applies to all networks and systems managed by {C} and any system that "
        f"connects to a {C}-managed network. This includes warehouse and office site networks, "
        "regional hub and colocation facilities, data center environments, cloud infrastructure "
        "(GCVE, GCE, and any other IaaS/PaaS tenants), and third-party managed systems connected "
        f"to {C} networks."
    ))
    add_para(doc, (
        "Networks and systems that cannot comply with this standard require a documented security "
        "exception approved by the Global Security Office."
    ))

    doc.add_page_break()

    # 4. Network Zones
    doc.add_heading("Network Zones", level=1)
    add_para(doc, (
        "The following network zones define the segmentation model. Not all zones will be present "
        "at every location; the applicable zones depend on the site type and function."
    ))

    doc.add_heading("Site-Level Zones", level=2)
    add_para(doc, "These zones apply to warehouse and office networks as defined in the Site Segmentation Standard:")
    add_bullet(doc, "Internet Zone \u2013 Public network segment. Only perimeter infrastructure devices (routers, firewalls) should have Layer 3 interfaces. Servers must not be directly exposed.")
    add_bullet(doc, "External DMZ \u2013 Systems requiring direct internet connectivity (client-facing applications for transient datasets, secure file transfer). Generally not expected at warehouse or office sites.")
    add_bullet(doc, "3rd Party Zone \u2013 Non-company-managed systems or devices. A separate zone must be created for each 3rd party unless coexistence is approved.")
    add_bullet(doc, f"{C} Operations Zone \u2013 Managed endpoints needing 3rd party access. Also includes local systems unable to run the standard security stack (IIoT, OT devices). If both types exist at a site, two separate Operations Zones are required.")
    add_bullet(doc, f"{C} Corporate Zone \u2013 All corporate devices and managed endpoints not requiring 3rd party access.")
    add_bullet(doc, "Internet Only Zone \u2013 Guest devices. Internet access only; no traffic to any other zone under any circumstances.")

    doc.add_heading("Hub / Data Center Zones", level=2)
    add_bullet(doc, "Internet Zone (DMZ) \u2013 Low trust zone (Security Level 50). All systems requiring direct internet connectivity. Implemented as a Layer 3 subnet.")
    add_bullet(doc, "Internet Support Zone \u2013 Medium trust zone. Supports the DMZ with middleware and backend services.")
    add_bullet(doc, "Infrastructure Zone \u2013 Secure zone (Security Level 55). Common infrastructure services (DNS, NTP, DHCP).")
    add_bullet(doc, "Server Zone \u2013 Internal zone (Security Level 100). Servers providing application services.")
    add_bullet(doc, "Management Zone \u2013 Internal zone (Security Level 100). IT environment management servers (SCCM, SolarWinds, etc.).")
    add_bullet(doc, "Security Zone \u2013 Internal zone (Security Level 100). Security service systems (EDR, SIEM, Tanium, Splunk, etc.).")
    add_bullet(doc, "Controller Zone \u2013 Systems owned/managed by 3rd parties. May include reduced-functionality OS that cannot be patched.")
    add_bullet(doc, "User Zone \u2013 End user compute devices: desktops, laptops, tablets, printers.")

    doc.add_heading("Cloud Zones", level=2)
    add_para(doc, "Cloud environments implement zone equivalence through platform-native controls:")
    add_bullet(doc, "GCVE \u2013 NSX distributed firewall provides micro-segmentation. Zones implemented as NSX security groups aligned to application tiers.")
    add_bullet(doc, "GCE \u2013 Google Cloud VPC firewall rules and network tags provide zone equivalence.")
    add_bullet(doc, f"Other Cloud (Azure, Oracle, AWS) \u2013 Equivalent native segmentation controls must be used. All cloud tenants must terminate via private link or VPN into a {C} datacenter or hub; private cloud connections must not terminate at a warehouse or office site.")

    doc.add_page_break()

    # 5. Security Levels
    doc.add_heading("Security Levels", level=1)
    add_para(doc, "Each zone is assigned a security level that determines trust relationships and default communication rules. Higher numbers indicate higher trust.")
    add_table(doc,
        ["Level", "Classification", "Zone Name", "Description"],
        [
            ["100", "Trusted", "Internal", f"{C} networks where users and internal services reside. Includes Production, Internal, End-users, Management, Security Zones."],
            ["60", "Trusted", "Isolated", f"{C} networks for QA, development, and lab devices. QA servers, Developer servers, DR-Test zones."],
            ["55", "Risky", "Secure", "Network between Trusted and DMZ. Common and Infrastructure zones."],
            ["50", "Risky", "DMZ", f"{C} controlled networks connected to external networks. DMZ servers."],
            ["30", "Risky", "Third Party", "IP addresses and devices belonging to a 3rd party with dedicated connection."],
            ["10", "Untrusted", "Guest Access", f"{C} networks for guests to access internet services only."],
            ["0", "Untrusted", "Internet", f"All public IP addresses not assigned to {C} or a Third Party."],
        ],
        col_widths=[0.6, 1.0, 1.0, 3.6],
    )

    doc.add_page_break()

    # 6. Inter-Zone Rules
    doc.add_heading("Inter-Zone Communications and Access Controls", level=1)
    add_para(doc, (
        "The security value of network segmentation comes from restricting communications between "
        "zones to only what is necessary. Unless explicitly permitted below, all inter-zone "
        "communication is denied by default."
    ))

    doc.add_heading("Default Rules", level=2)
    add_numbered(doc, "All inter-zone traffic, both inbound and outbound, must be denied by default.")
    add_numbered(doc, "Firewall rules and ACLs must never contain a source address or destination address/protocol/port of \u2018any\u2019 unless the rule is for a proxy server managing internet access.")
    add_numbered(doc, "The concept of minimum access must be enforced: every rule allows only the minimum traffic required to support a specific business requirement.")
    add_numbered(doc, "All rules must have a documented business justification and a defined expiry review period.")
    add_numbered(doc, "Zone/segment separation must be performed by a firewall, with configurations governed by Firemon or equivalent policy management.")

    doc.add_heading("Internet Zone Rules", level=2)
    add_bullet(doc, "All connectivity from the Internet must terminate in the Internet Zone (DMZ). No other zones may accept external connections directly.")
    add_bullet(doc, "Systems in the Internet Zone may access specific systems in the Internet Support Zone as required. The reverse is not permitted.")
    add_bullet(doc, "The Infrastructure Zone may access the Internet Zone (and vice versa) only for infrastructure services such as DNS and NTP, limited to specific protocols, ports, and devices.")

    doc.add_heading("Security and Management Zone Rules", level=2)
    add_bullet(doc, "The Security and Management Zones may access systems in all other zones as required for security monitoring, patching, and management.")
    add_bullet(doc, "Other zones may not initiate connections to the Security and Management Zones except via designated jump servers and only by authorized personnel.")

    doc.add_heading("User Zone Rules", level=2)
    add_bullet(doc, "The User Zone may access the Internet Zone only as required for internet proxy services.")
    add_bullet(doc, "The User Zone may access the Infrastructure Zone only to receive infrastructure services.")
    add_bullet(doc, "The User Zone may access the Server Zone only as required to access internal applications.")
    add_bullet(doc, "The User Zone may access the Security and Management Zones only via jump server and only by authorized personnel.")

    doc.add_heading("Infrastructure Zone", level=2)
    add_bullet(doc, "All zones may access the Infrastructure Zone as needed to receive infrastructure services (DNS, NTP, DHCP, etc.).")
    add_bullet(doc, "The Infrastructure Zone must not be used as a transit zone for traffic between other zones.")

    doc.add_heading("Site-to-Site Communication", level=2)
    add_para(doc, (
        "Site-to-site and site-to-datacenter communications are not permitted unless required to "
        "support a documented business process. A valid business justification must be provided "
        "to the GSO for each connection, limited to specific protocols, ports, and devices."
    ))
    add_para(doc, (
        "The goal is to eliminate site-to-site exceptions over time by routing cross-site "
        "application access through the cloud or datacenter."
    ))

    doc.add_page_break()

    # 7. Tiered Enforcement
    doc.add_heading("Tiered Enforcement Model", level=1)
    add_para(doc, "This standard adopts a tiered enforcement model to ensure security controls are applied at the most effective point without duplication:")
    add_table(doc,
        ["Tier", "Location", "Inspection Responsibility", "Controls"],
        [
            ["1", "Site", "Zone separation and north/south access control via stateful ACLs", "Stateful Packet Inspection"],
            ["2", "Hub (Equinix)", "Deep inspection, threat detection, anti-malware, and service-based rules (transitional)", "DPI, IPS/IDS, Anti-Malware"],
            ["3", "Cloud (GCVE/GCE)", "Application-level micro-segmentation at the resource", "NSX DFW / VPC FW rules"],
            ["0", "Cloud Proxy", "Web and SaaS traffic inspection, URL filtering, cloud DLP", "SWG"],
        ],
        col_widths=[0.5, 1.2, 2.8, 1.8],
    )
    add_para(doc, (
        "Controls must not be duplicated between tiers unless a specific compliance or contractual "
        "requirement mandates it. Any such duplication must be documented with a business "
        "justification and approved by the GSO."
    ))
    add_para(doc, (
        f"Refer to the {C} Network Security Strategy (DOC-IT-IS-NSS-001) for the full "
        "enforcement point matrix and traffic flow walkthroughs."
    ))
    add_para(doc, (
        "This standard defines security requirements independently of enforcement technology. "
        "The requirements remain valid regardless of whether inspection is performed by "
        "on-premise firewalls, cloud-delivered services, or a combination of both."
    ), size=11)

    doc.add_page_break()

    # 8. External Access
    doc.add_heading("External Access to Systems", level=1)
    add_para(doc, (
        "External access to network zones must be provided through the corporate VPN solution "
        "using multi-factor authentication, or through an approved Zero Trust Network Access "
        "(ZTNA) solution where deployed."
    ))
    add_para(doc, "Remote access must only provide access to the applications, tools, and systems required to perform job duties. Only authorized administrative and support personnel may access jump servers.")
    add_para(doc, "Third-party remote access must terminate at a company-controlled firewall. Where the third party cannot terminate at the firewall, a separate 3rd Party Network must be created by exception to host the VPN termination point.")

    # 9. Exception Severity
    doc.add_heading("Exception Severity Classification", level=1)
    add_para(doc, "Deviations from this standard are classified by severity:")
    add_severity_table(doc,
        ["Severity", "Trigger", "Action Required"],
        [
            ["Critical",
             "Traffic from Internet to Trusted, Secure, Isolated, or Third Party zones",
             "Security exception required with business justification. Exception active for defined time frame only."],
            ["High",
             "Source/destination of \u2018any\u2019 in or out of Secure, Isolated, DMZ, or Third Party. All traffic from Internet to DMZ.",
             "Exception must be approved by Information Security. Active for pre-defined time frame."],
            ["Medium",
             "Source/destination of \u2018any\u2019 in or out of Internal Zone. Traffic DMZ to Internal or Internet to DMZ. Rules unused for 240 days.",
             "Business as usual with highly permissive rules flagged. Warning for unused rules before 365-day limit."],
            ["Low",
             "Internal traffic requiring ACL policy guidelines. Traffic to/from Guest zone (except Guest to Internet). Inter-zone traffic with source/destination \u2018any\u2019.",
             "Business as usual. Used to notify security of rules requiring review."],
        ],
        col_widths=[0.8, 3.0, 2.7],
    )

    # 10. Compliance
    doc.add_heading("Compliance and Review", level=1)
    add_para(doc, (
        "This standard must be reviewed annually or following any significant change to the "
        "network architecture (e.g., new cloud provider adoption, major site rollout methodology "
        "change, or changes to the enforcement technology model)."
    ))
    add_para(doc, (
        "All firewall rules must be audited quarterly against this standard using Firemon or "
        "equivalent. Rules that do not comply and do not have an approved exception must be "
        "remediated within 30 days of identification."
    ))
    add_para(doc, (
        "Any rule unused for 365 days must be removed unless a business justification for "
        "retention is approved by the GSO."
    ))

    setup_header_footer(doc, f"{C} Network Segmentation Standard v2.0")
    return doc


# ============================================================
# DOCUMENT 3: INSPECTION & THREAT PROTECTION STANDARD
# ============================================================

def create_inspection_standard():
    doc = new_doc()

    add_title_page(
        doc,
        "Inspection & Threat Protection Standard",
        doc_id="DOC-IT-IS-ITP-001",
        version="1.0 DRAFT",
        owner="Global Security Office",
    )

    # 1. Purpose
    doc.add_heading("Purpose", level=1)
    add_para(doc, (
        f"This standard defines the requirements for network traffic inspection and threat "
        f"protection across the {C} network. It specifies which inspection technologies must be "
        "deployed, where in the network path they must be placed, the inspection profiles to be "
        "used, and the operational requirements for maintaining and tuning these controls."
    ))
    add_para(doc, "This document addresses the gap in legacy standards where IPS/IDS, deep packet inspection, and anti-malware requirements were not formally defined.")

    # 2. Scope
    doc.add_heading("Scope", level=1)
    add_para(doc, (
        f"This standard applies to all traffic inspection and threat protection technologies "
        f"deployed on {C} networks, including Intrusion Prevention/Detection Systems (IPS/IDS), "
        "Deep Packet Inspection (DPI) engines, anti-malware and file scanning engines, SSL/TLS "
        "inspection capabilities, Web Application Firewalls (WAF) where deployed, and "
        "cloud-delivered inspection via the cloud proxy (SWG)."
    ))

    doc.add_page_break()

    # 3. Inspection Technologies
    doc.add_heading("Inspection Technologies Defined", level=1)

    doc.add_heading("Stateful Packet Inspection (SPI)", level=2)
    add_para(doc, "Stateful packet inspection tracks the state of network connections (TCP sessions, UDP streams) and makes forwarding decisions based on connection state, source/destination IP, and port numbers. SPI does not examine the payload content of packets.")
    add_bold_para(doc, "Use case: ", "Access control and zone enforcement at site perimeters where performance is critical and the traffic will be inspected deeper further along the path.")
    add_bold_para(doc, "Equipment requirements: ", "Any Layer 3 routing device with ACL capability, or an SD-WAN appliance with stateful firewall features. Does not require advanced threat protection licensing.")

    doc.add_heading("Deep Packet Inspection (DPI)", level=2)
    add_para(doc, "Deep packet inspection examines the full content of network packets beyond the header, including the application layer payload. DPI can identify applications regardless of port number, detect protocol anomalies, and feed content to IPS and anti-malware engines.")
    add_bold_para(doc, "Use case: ", "Primary threat inspection point at hub locations where all site-to-cloud traffic transits. Provides the foundation for IPS and anti-malware detection.")
    add_bold_para(doc, "Equipment requirements: ", "Next-generation firewall with DPI engine and sufficient throughput for aggregate traffic from all connected sites. Throughput must be rated for DPI-enabled traffic, not just stateful inspection.")

    doc.add_heading("Intrusion Prevention / Detection System (IPS/IDS)", level=2)
    add_para(doc, "IPS examines traffic against a database of known attack signatures and behavioral anomalies. In prevention mode (IPS), detected threats are blocked inline. In detection mode (IDS), threats are logged and alerted but not blocked.")
    add_bold_para(doc, "Use case: ", "Threat detection and prevention at network transit points. Signatures must cover exploit attempts, command-and-control traffic, malware indicators, and protocol abuse.")
    add_bold_para(doc, "Mode of operation: ", "IPS (inline prevention) is the default mode at hub firewalls. IDS (detection only) may be used during initial deployment or tuning periods, but must transition to IPS within 90 days.")

    doc.add_heading("Anti-Malware Scanning", level=2)
    add_para(doc, "Network-based anti-malware scanning inspects file transfers traversing the network for known malware signatures and suspicious file characteristics. This complements endpoint-based anti-malware.")
    add_bold_para(doc, "Use case: ", "Scanning files in transit at the hub layer. Particularly important for traffic entering from 3rd party networks or the internet.")

    doc.add_heading("SSL/TLS Inspection", level=2)
    add_para(doc, "SSL/TLS inspection decrypts encrypted traffic for DPI/IPS analysis and re-encrypts it before forwarding. Without SSL inspection, encrypted traffic bypasses DPI and IPS engines.")
    add_bold_para(doc, "Use case: ", "Applied at the cloud proxy layer for internet-bound traffic. Optional at the hub layer for internal encrypted traffic where compliance requirements mandate it.")

    doc.add_heading("Web Application Firewall (WAF)", level=2)
    add_para(doc, "WAF provides application-layer protection for web-facing services, protecting against OWASP Top 10 vulnerabilities, SQL injection, cross-site scripting, and other web application attacks.")
    add_bold_para(doc, "Use case: ", "Required for any company-hosted web application exposed to the internet or to untrusted zones. Can be cloud-delivered or on-premise.")
    add_bold_para(doc, "Note: ", "WAF requirements should be addressed in a separate standard or as an addendum to this document. WAF is listed here for completeness.")

    doc.add_page_break()

    # 4. Placement Requirements
    doc.add_heading("Inspection Placement Requirements", level=1)
    add_para(doc, "The following table defines the mandatory and optional inspection capabilities at each enforcement point. This is the authoritative reference to prevent duplication of inspection.")
    add_table(doc,
        ["Enforcement Point", "SPI", "DPI / IPS", "Anti-Malware", "SSL Inspection"],
        [
            ["Site Firewall / SD-WAN", "MANDATORY", "NOT REQUIRED", "NOT REQUIRED", "NOT REQUIRED"],
            ["Hub Firewall (Equinix)", "MANDATORY", "MANDATORY", "MANDATORY", "OPTIONAL"],
            ["Cloud Edge (NSX / VPC)", "N/A", "NOT REQUIRED", "NOT REQUIRED", "NOT REQUIRED"],
            ["Cloud Proxy (SWG)", "N/A", "MANDATORY (web)", "MANDATORY (web)", "MANDATORY"],
            ["TSG Datacenter", "MANDATORY", "MANDATORY", "MANDATORY", "OPTIONAL"],
        ],
        col_widths=[1.6, 1.0, 1.1, 1.0, 1.0],
    )
    add_para(doc, "NOT REQUIRED means the control is explicitly excluded at that point to avoid duplication. If a specific compliance requirement mandates additional inspection at a point marked NOT REQUIRED, a documented exception must be raised and approved by the GSO.")

    doc.add_page_break()

    # 5. IPS Profile Requirements
    doc.add_heading("IPS Profile Requirements", level=1)

    doc.add_heading("Signature Management", level=2)
    add_bullet(doc, "IPS signatures must be updated at least weekly. Critical out-of-band updates for actively exploited zero-day vulnerabilities must be applied within 24 hours of vendor release.")
    add_bullet(doc, "The signature database must include coverage for exploit attempts against known CVEs, command-and-control (C2) communication patterns, DNS tunneling and exfiltration techniques, protocol anomalies and abuse, and brute force and credential stuffing patterns.")
    add_bullet(doc, "Signatures should be sourced from the firewall vendor\u2019s threat intelligence feed. Where available, custom signatures should be developed for threats targeting the logistics sector.")

    doc.add_heading("Inspection Profiles", level=2)
    add_para(doc, "Most firewall vendors provide pre-built inspection profiles at low, medium, and high sensitivity. The following approach should be used:")
    add_table(doc,
        ["Profile", "Enforcement Point", "Behaviour", "Performance Impact"],
        [
            ["High", "Hub FW (site-to-cloud)", "Maximum detection sensitivity. All signatures enabled. May generate more false positives during tuning.", "Higher latency (10\u201320ms). Requires adequately sized hardware."],
            ["Medium", "Hub FW (inter-DC)", "Balanced detection. High-confidence signatures enabled. Anomaly detection active.", "Moderate latency (5\u201310ms)."],
            ["Low", "Not recommended", "Basic detection only. Many signatures disabled. Suitable only for initial deployment / tuning.", "Minimal latency."],
        ],
        col_widths=[0.8, 1.3, 2.5, 1.7],
    )
    add_para(doc, (
        "The default IPS profile at hub firewalls must be High for all traffic originating from "
        "sites destined for cloud/datacenter environments. The profile may be tuned to exclude "
        "specific signatures only where documented false positives have been identified and "
        "validated by the security operations team."
    ))

    doc.add_heading("Performance Considerations", level=2)
    add_bullet(doc, "Hub firewall hardware must be sized for DPI/IPS throughput at the aggregate bandwidth of all connected sites, with at least 30% headroom for traffic growth.")
    add_bullet(doc, "Latency through the hub inspection point should be monitored. If sustained inspection latency exceeds 20ms for latency-sensitive applications (WMS, voice, video), the security operations team and network engineering must collaborate on optimization.")
    add_bullet(doc, "Application-specific bypass rules (where a specific application flow bypasses DPI/IPS) must be documented as security exceptions and approved by the GSO. They should be reviewed quarterly.")

    doc.add_page_break()

    # 6. Inspection Position
    doc.add_heading("Inspection Position in Traffic Flow", level=1)
    add_para(doc, "The following principles determine where inspection occurs in the path:")
    add_bullet(doc, "Internet ingress traffic: Inspected by the cloud proxy for web traffic, or by the hub/DC firewall for VPN or other non-web ingress.")
    add_bullet(doc, "Site-to-cloud traffic: Inspected at the hub (Tier 2) once. Not re-inspected at the cloud entry point unless a specific compliance mandate requires it.")
    add_bullet(doc, "Cloud-to-site traffic: Inspected at the hub (Tier 2) once on the return path.")
    add_bullet(doc, "Intra-cloud traffic (e.g., GCVE VM to GCVE VM): Governed by NSX micro-segmentation rules. No DPI/IPS unless enabled for a specific compliance scenario.")
    add_bullet(doc, "Site-to-site traffic (exception-based): Inspected at the hub (Tier 2) as the traffic transits the hub.")
    add_bullet(doc, "Internet egress traffic: Inspected by the cloud proxy. Does not transit the hub.")
    add_para(doc, "The key principle is single inspection per flow. Traffic should not be inspected by DPI/IPS at more than one point on its path unless there is a documented, approved reason for doing so.")

    # 7. Operational Requirements
    doc.add_heading("Operational Requirements", level=1)

    doc.add_heading("Monitoring and Alerting", level=2)
    add_bullet(doc, "All IPS events must be forwarded to the SIEM (Splunk or equivalent) in real-time.")
    add_bullet(doc, "Critical and high-severity IPS alerts must generate an incident ticket within 15 minutes.")
    add_bullet(doc, "The security operations team must review IPS event trends weekly and produce a monthly report for the GSO.")

    doc.add_heading("Tuning and False Positive Management", level=2)
    add_bullet(doc, "False positives must be documented and tracked. A signature must not be disabled without a documented tuning request approved by the security operations team lead.")
    add_bullet(doc, "Tuning changes must be tested in detection mode (IDS) before being applied in prevention mode (IPS).")
    add_bullet(doc, "A quarterly review of all disabled signatures must be conducted to determine if they can be re-enabled.")

    doc.add_heading("Capacity Planning", level=2)
    add_bullet(doc, "Hub firewall utilization (CPU, memory, session count) must be monitored and alerted at 70% sustained utilization.")
    add_bullet(doc, "When utilization exceeds 80% sustained, a capacity planning exercise must be initiated.")
    add_bullet(doc, "Any new site onboarding or cloud workload migration expected to increase hub traffic by more than 10% must include a capacity impact assessment.")

    doc.add_page_break()

    # 8. Technology Neutrality (replaces the old full SASE section)
    doc.add_heading("Technology Neutrality", level=1)
    add_para(doc, (
        "This standard defines inspection requirements in terms of what must be inspected and "
        "at what depth, rather than mandating specific products or deployment models. The "
        "requirements are valid whether inspection is performed by on-premise next-generation "
        "firewalls, cloud-delivered security services, or a combination of both."
    ))
    add_para(doc, (
        "If the organization transitions to a different enforcement technology in the future "
        "(for example, shifting DPI/IPS from hub-based firewalls to a cloud-delivered inspection "
        "service), the inspection requirements defined in this standard remain applicable. Only "
        "the placement table in Section 4 would need to be updated to reflect the new enforcement "
        "points."
    ))

    # 9. Compliance
    doc.add_heading("Compliance and Review", level=1)
    add_para(doc, (
        "This standard must be reviewed annually or when significant changes occur to the "
        "inspection infrastructure (new firewall platforms, changes to enforcement technology, "
        "new compliance requirements)."
    ))
    add_para(doc, (
        "Compliance will be assessed through quarterly firewall configuration audits via Firemon, "
        "monthly IPS event and tuning reports, and annual penetration testing that validates "
        "inspection effectiveness."
    ))

    setup_header_footer(doc, f"{C} Inspection & Threat Protection Standard")
    return doc


# ============================================================
# MAIN
# ============================================================

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    docs = [
        (create_strategy, f"{COMPANY_NAME}_Network_Security_Strategy.docx"),
        (create_netseg_standard, f"{COMPANY_NAME}_Network_Segmentation_Standard_v2.docx"),
        (create_inspection_standard, f"{COMPANY_NAME}_Inspection_Threat_Protection_Standard.docx"),
    ]

    for creator, filename in docs:
        filepath = os.path.join(OUTPUT_DIR, filename)
        print(f"Generating {filename}...")
        doc = creator()
        doc.save(filepath)
        print(f"  \u2713 Saved to {filepath}")

    print(f"\nDone! {len(docs)} documents generated in '{OUTPUT_DIR}/'")
    print(f"\nCompany name used: '{COMPANY_NAME}'")
    print("To customize, edit COMPANY_NAME at the top of this script.")


if __name__ == "__main__":
    main()
