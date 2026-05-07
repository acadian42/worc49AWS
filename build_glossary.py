"""
LLC Network Security Strategy — Glossary Builder
================================================
Generates a Word document containing a styled glossary table that can be
copy/pasted into the Network Security Strategy document.

Color scheme matches the existing LLC document tables:
  - Header row:        warm orange (#C65911) with white bold text
  - Body rows:         alternating white and light-orange tint (#FBE5D6)
  - Borders:           light grey

Run:
    python build_glossary.py
Output:
    LLC_Strategy_Glossary.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ----------------------------------------------------------------------
# Style constants — change these if the brand orange is slightly different
# ----------------------------------------------------------------------
HEADER_FILL   = "C65911"   # warm orange, matches existing LLC tables
ALT_ROW_FILL  = "FBE5D6"   # light tint of the header for alternating rows
BORDER_COLOR  = "BFBFBF"   # light grey
FONT_NAME     = "Arial"


# ----------------------------------------------------------------------
# Glossary content
# ----------------------------------------------------------------------
# NOTE: "AV LAN" is included as a placeholder — the term came up verbally
# but is not defined in any of the source documents. Fill in your org's
# working definition before publishing.
# ----------------------------------------------------------------------
GLOSSARY = [
    ("Anti-Malware",
     "An inspection capability that scans network traffic and files for known "
     "malicious code using signature databases, heuristic analysis, and/or "
     "sandbox detonation. Typically deployed alongside DPI/IPS at the hub "
     "and at the SSE layer for internet-bound web traffic."),

    ("AV LAN",
     "[TBD — placeholder. Define the LLC working meaning. Commonly refers to a "
     "dedicated network segment for audio/visual and conference-room endpoints "
     "(displays, room controllers, video bars) that cannot host the standard "
     "security stack and are isolated from the Corporate Zone.]"),

    ("Blast Radius",
     "The scope of systems and data that can be reached or affected from a "
     "single compromised host. Network segmentation reduces blast radius by "
     "limiting lateral movement options after an initial compromise."),

    ("CASB (Cloud Access Security Broker)",
     "A control point that sits between users and SaaS applications to enforce "
     "policy on access, data movement, and risky behaviour. Delivered at the "
     "SSE layer (Tier 0) in the LLC model."),

    ("DFW (Distributed Firewall)",
     "A firewall enforced at the hypervisor or workload level rather than at a "
     "network choke point. Enables micro-segmentation by applying rules between "
     "individual VMs regardless of their physical location."),

    ("DLP (Data Loss Prevention)",
     "Controls that detect and block sensitive data from leaving the "
     "environment in unauthorised ways. Delivered primarily at the SSE layer "
     "(Tier 0) for cloud and web egress."),

    ("DNS Security",
     "Protection of DNS resolution against malicious domains, DNS tunnelling, "
     "and exfiltration. Typically delivered via secure resolvers at the hub "
     "and SSE layers."),

    ("DPI (Deep Packet Inspection)",
     "Inspection of the full contents of network packets — including payload — "
     "to identify applications, threats, and policy violations. Effective on "
     "encrypted traffic only when combined with TLS decryption."),

    ("East-West Traffic",
     "Network traffic flowing between workloads inside the same environment "
     "(e.g. VM to VM within GCE, or server to server within a data centre). "
     "Distinct from north-south traffic, which crosses an environment boundary."),

    ("Enforcement Point",
     "A network location where one or more security controls are applied. "
     "The strategy defines four enforcement points: Tier 0 (SSE), Tier 1 "
     "(Site), Tier 2 (Hub), and Tier 3 (Cloud Workload)."),

    ("GCE (Google Compute Engine)",
     "Google Cloud's IaaS compute service. In the LLC model, GCE workloads "
     "are protected by VPC firewall rules and cloud-native NGFW capabilities "
     "as the Tier 3 enforcement point."),

    ("GCVE (Google Cloud VMware Engine)",
     "A managed VMware environment running in Google Cloud. Uses NSX for "
     "micro-segmentation as the Tier 3 enforcement point in the LLC model."),

    ("IDPS (Intrusion Detection and Prevention System)",
     "Combined term for IPS and IDS capabilities — detection and active "
     "blocking of network-based attacks based on signatures, behavioural "
     "rules, and threat intelligence."),

    ("IPS / IDS (Intrusion Prevention / Detection System)",
     "Inspection technology that identifies attempted exploits, malware "
     "command-and-control, and other malicious patterns in traffic. IPS "
     "blocks in line; IDS detects and alerts only."),

    ("Jump Server",
     "A hardened intermediary host used by authorised personnel to reach "
     "Security and Management zones. The only sanctioned path into those "
     "zones from User and other lower-trust zones."),

    ("Lateral Movement",
     "An attacker's progression from an initial foothold to additional "
     "systems within the environment. Segmentation, micro-segmentation, "
     "and east-west inspection are the primary controls against it."),

    ("Micro-Segmentation",
     "Fine-grained policy enforcement between individual workloads — typically "
     "VM-to-VM or container-to-container — using a distributed firewall (NSX) "
     "or cloud-native firewall (VPC firewall rules). Operates at Tier 3."),

    ("NGFW (Next-Generation Firewall)",
     "A firewall combining stateful packet filtering with application "
     "awareness, IPS, anti-malware, and (where licensed) TLS inspection. "
     "Deployed at the hub (Tier 2) and increasingly cloud-native at Tier 3."),

    ("North-South Traffic",
     "Network traffic crossing an environment boundary — for example, "
     "internet to data centre, site to cloud, or user to SaaS. Distinct "
     "from east-west traffic."),

    ("NSX (VMware NSX)",
     "VMware's network virtualisation and security platform. Provides the "
     "distributed firewall used for micro-segmentation in GCVE."),

    ("PKI (Public Key Infrastructure)",
     "The framework of certificate authorities, certificates, and trust "
     "relationships used to issue and validate digital certificates. "
     "Enterprise PKI is required for inspection devices to perform TLS "
     "decryption without breaking client trust."),

    ("S2S (Site-to-Site)",
     "Direct network connectivity between two LLC sites, typically over "
     "VPN or SD-WAN. Permitted only by exception with documented business "
     "justification; the strategic direction is to route cross-site flows "
     "via the cloud or data centre instead."),

    ("SASE (Secure Access Service Edge)",
     "A converged architecture delivering both networking (SD-WAN) and "
     "security (SSE) as a cloud service. The future-state target for the "
     "LLC inspection model."),

    ("SD-WAN (Software-Defined WAN)",
     "WAN architecture that uses software policy to steer traffic across "
     "multiple transport links. Provides the Tier 1 site firewall function "
     "in many LLC sites."),

    ("Service-Based Application Rules",
     "Firewall rules that allow traffic based on the application or service "
     "identified (e.g. Salesforce, M365) rather than purely on port and "
     "protocol. Primary at Tier 3 (cloud); transitional at Tier 2 (hub)."),

    ("SPI (Stateful Packet Inspection)",
     "Layer-3/4 firewalling that tracks connection state to allow only "
     "legitimate response traffic. The baseline control at every "
     "enforcement point and the primary function at Tier 1 (site)."),

    ("SSE (Security Service Edge)",
     "Cloud-delivered security services — typically SWG, CASB, ZTNA, and "
     "cloud DLP. Tier 0 in the LLC enforcement model. Zscaler is the "
     "current LLC SSE platform."),

    ("SSL / TLS Inspection",
     "Decryption of encrypted traffic by an inspection device for DPI/IPS/"
     "anti-malware analysis, followed by re-encryption before forwarding. "
     "Requires enterprise PKI trust on endpoints. Without it, encrypted "
     "payloads are not visible to inspection engines."),

    ("SWG (Secure Web Gateway)",
     "URL filtering, malware scanning, and policy enforcement for user "
     "internet-bound web traffic. Delivered at the SSE layer (Tier 0)."),

    ("Tier 0 — SSE",
     "Cloud-delivered inspection layer (Zscaler) responsible for web and "
     "SaaS traffic, URL filtering, cloud DLP, and (future) ZTNA."),

    ("Tier 1 — Site",
     "Site firewall / SD-WAN appliance. Performs stateful packet inspection "
     "and zone separation. No DPI by default; narrow exceptions for "
     "site-to-site traffic and visiting third-party flows."),

    ("Tier 2 — Hub",
     "Hub firewalls in colocation facilities (e.g. Equinix) and the legacy "
     "data centre. Primary enforcement point for DPI, IPS/IDS, and "
     "anti-malware on WAN traffic. Service-based app rules treated as "
     "transitional pending full cloud migration."),

    ("Tier 3 — Cloud Workload",
     "Workload-level enforcement in cloud environments. NSX distributed "
     "firewall in GCVE; VPC firewall rules and cloud-native NGFW in GCE. "
     "Primary enforcement point for micro-segmentation."),

    ("URL / Web Filtering",
     "Category-based and reputation-based blocking of web destinations for "
     "user-initiated browsing. Primary at the SSE layer (Tier 0)."),

    ("VPC (Virtual Private Cloud)",
     "An isolated network within a public cloud provider, with native "
     "firewall rules used for zone equivalence and segmentation in GCE."),

    ("WAF (Web Application Firewall)",
     "Application-layer protection for web-facing services against OWASP "
     "Top 10 and similar attacks. Out of scope for this strategy — to be "
     "addressed in a separate standard or addendum."),

    ("Zone Separation (VLAN)",
     "Layer-2/3 isolation of network segments using VLANs and routed "
     "subnets, with inter-zone traffic forced through a firewall. The "
     "primary segmentation mechanism at Tier 1 sites."),

    ("ZTNA (Zero Trust Network Access)",
     "Identity- and posture-aware access to specific applications, "
     "replacing traditional broad VPN access. Delivered at the SSE layer; "
     "on the LLC roadmap as Zscaler ZTNA matures."),
]


# ----------------------------------------------------------------------
# Helpers — direct XML for things python-docx doesn't expose cleanly
# ----------------------------------------------------------------------
def shade_cell(cell, hex_fill: str) -> None:
    """Apply a solid background colour to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_borders(cell, hex_color: str = BORDER_COLOR, size: str = "4") -> None:
    """Apply consistent thin borders to a single cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), hex_color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def style_run(run, *, bold: bool = False, color_rgb: RGBColor | None = None,
              size_pt: int = 10) -> None:
    """Apply font, size, weight, and colour to a run."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_rgb is not None:
        run.font.color.rgb = color_rgb
    # Force the East-Asian font slot to also use Arial — avoids fallback
    # rendering in some Word versions.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)


# ----------------------------------------------------------------------
# Build document
# ----------------------------------------------------------------------
def build(output_path: str) -> None:
    doc = Document()

    # Page margins — 1 inch all sides, US Letter
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Glossary of Terms")
    style_run(title_run, bold=True, size_pt=18)

    # Intro paragraph
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "The following terms are used throughout this strategy. Definitions "
        "are scoped to their meaning within the LLC Network Security Strategy "
        "and may be more specific than general industry usage."
    )
    style_run(intro_run, size_pt=10)

    doc.add_paragraph()  # spacer

    # Sort glossary alphabetically
    entries = sorted(GLOSSARY, key=lambda x: x[0].lower())

    # Build the table: 1 header row + N body rows, 2 columns
    table = doc.add_table(rows=1 + len(entries), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Column widths — Term column narrower, Definition wider
    col_widths = (Inches(2.0), Inches(4.5))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]

    # Header row
    header = table.rows[0]
    header_cells = header.cells
    header_labels = ("Term", "Definition")
    for cell, label in zip(header_cells, header_labels):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, HEADER_FILL)
        set_cell_borders(cell)
        # Replace default empty paragraph
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(label)
        style_run(run, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF), size_pt=11)

    # Body rows
    for i, (term, definition) in enumerate(entries):
        row = table.rows[i + 1]
        term_cell, def_cell = row.cells

        # Alternating row shading: shade odd rows (i.e. every second body row)
        if i % 2 == 1:
            shade_cell(term_cell, ALT_ROW_FILL)
            shade_cell(def_cell, ALT_ROW_FILL)

        for cell in (term_cell, def_cell):
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Term cell — bold
        p1 = term_cell.paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after = Pt(3)
        r1 = p1.add_run(term)
        style_run(r1, bold=True, size_pt=10)

        # Definition cell
        p2 = def_cell.paragraphs[0]
        p2.paragraph_format.space_before = Pt(3)
        p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run(definition)
        style_run(r2, size_pt=10)

    doc.save(output_path)
    print(f"Wrote: {output_path}  ({len(entries)} entries)")


if __name__ == "__main__":
    build("LLC_Strategy_Glossary.docx")
