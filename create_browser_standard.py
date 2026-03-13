"""
Browser Standardisation Standard - Document Generator
Requirements: pip install python-docx
Usage:        python create_browser_standard.py
Output:       Browser_Standardisation_Standard.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colours ──────────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x4E, 0x79)
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE = "D6E4F0"
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY       = RGBColor(0x55, 0x55, 0x55)
GREEN_FG   = RGBColor(0x1D, 0x6A, 0x2E)
GREEN_BG   = "E8F5E9"
RED_FG     = RGBColor(0x8B, 0x00, 0x00)
RED_BG     = "FDECEA"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour):
    """Fill a table cell with a solid background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def set_cell_borders(cell):
    """Apply a thin grey border to all four sides of a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Set internal cell padding (in twentieths of a point)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def add_bottom_border_to_para(para, colour="2E75B6", size="12"):
    """Draw a coloured bottom border under a paragraph."""
    pPr    = para._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_run(para, text, bold=False, italic=False,
            colour=None, size_pt=11, font="Arial"):
    run          = para.add_run(text)
    run.bold     = bold
    run.italic   = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    if colour:
        run.font.color.rgb = colour
    return run


# ── Section builders ──────────────────────────────────────────────────────────

def add_heading1(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after  = Pt(6)
    add_bottom_border_to_para(para)
    add_run(para, text, bold=True, colour=DARK_BLUE, size_pt=14)
    return para


def add_heading2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    add_run(para, text, bold=True, colour=MID_BLUE, size_pt=12)
    return para


def add_body(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    add_run(para, text, size_pt=11)
    return para


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    add_run(para, text, size_pt=11)
    return para


def add_info_table(doc, rows):
    """Two-column key/value metadata table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    col_widths  = [Inches(1.8), Inches(5.0)]

    for i, (label, value) in enumerate(rows):
        row   = table.rows[i]
        cells = row.cells

        # Label cell
        set_cell_bg(cells[0], LIGHT_BLUE)
        set_cell_borders(cells[0])
        set_cell_margins(cells[0])
        cells[0].width = col_widths[0]
        p0 = cells[0].paragraphs[0]
        add_run(p0, label, bold=True, colour=DARK_BLUE, size_pt=10)

        # Value cell
        set_cell_borders(cells[1])
        set_cell_margins(cells[1])
        cells[1].width = col_widths[1]
        p1 = cells[1].paragraphs[0]
        add_run(p1, value, size_pt=10)

    return table


def add_browser_table(doc, browsers):
    """
    Browser approval status table.
    browsers = list of (name, status, notes, is_approved)
    """
    headers    = ["Browser", "Status", "Notes"]
    col_widths = [Inches(1.8), Inches(1.3), Inches(3.7)]

    table = doc.add_table(rows=1 + len(browsers), cols=3)
    table.style = "Table Grid"

    # Header row
    hrow = table.rows[0]
    for j, hdr in enumerate(headers):
        c = hrow.cells[j]
        set_cell_bg(c, "1F4E79")
        set_cell_borders(c)
        set_cell_margins(c)
        c.width = col_widths[j]
        p = c.paragraphs[0]
        if j == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, hdr, bold=True, colour=WHITE, size_pt=10)

    # Data rows
    for i, (browser, status, notes, approved) in enumerate(browsers):
        row    = table.rows[i + 1]
        cells  = row.cells
        fg     = GREEN_FG if approved else RED_FG
        bg     = GREEN_BG if approved else RED_BG

        # Browser name
        set_cell_borders(cells[0])
        set_cell_margins(cells[0])
        cells[0].width = col_widths[0]
        if i % 2 == 0:
            set_cell_bg(cells[0], "F5F9FD")
        add_run(cells[0].paragraphs[0], browser, bold=True, size_pt=10)

        # Status
        set_cell_bg(cells[1], bg)
        set_cell_borders(cells[1])
        set_cell_margins(cells[1])
        cells[1].width = col_widths[1]
        p1 = cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p1, status, bold=True, colour=fg, size_pt=10)

        # Notes
        set_cell_borders(cells[2])
        set_cell_margins(cells[2])
        cells[2].width = col_widths[2]
        if i % 2 == 0:
            set_cell_bg(cells[2], "F5F9FD")
        add_run(cells[2].paragraphs[0], notes, size_pt=10)

    return table


def add_roles_table(doc, rows):
    """Two-column roles & responsibilities table."""
    col_widths = [Inches(2.1), Inches(4.7)]
    table      = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"

    # Header
    hrow = table.rows[0]
    for j, hdr in enumerate(["Role", "Responsibility"]):
        c = hrow.cells[j]
        set_cell_bg(c, "1F4E79")
        set_cell_borders(c)
        set_cell_margins(c)
        c.width = col_widths[j]
        add_run(c.paragraphs[0], hdr, bold=True, colour=WHITE, size_pt=10)

    # Data rows
    for i, (role, resp) in enumerate(rows):
        row   = table.rows[i + 1]
        cells = row.cells
        bg    = LIGHT_BLUE if i % 2 == 0 else None

        for j, (cell, txt, bold) in enumerate(
                zip(cells, [role, resp], [True, False])):
            if bg:
                set_cell_bg(cell, LIGHT_BLUE)
            set_cell_borders(cell)
            set_cell_margins(cell)
            cell.width = col_widths[j]
            add_run(cell.paragraphs[0], txt, bold=bold, size_pt=10)

    return table


# ── Main document builder ─────────────────────────────────────────────────────

def build_document(output_path="Browser_Standardisation_Standard.docx"):
    doc = Document()

    # Page margins (A4)
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)

    # ── Title block ───────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after  = Pt(4)
    add_run(title, "WEB BROWSER STANDARDISATION STANDARD",
            bold=True, colour=DARK_BLUE, size_pt=22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(16)
    add_run(sub, "IT Security Policy", colour=GREY, size_pt=13)

    # ── Metadata table ────────────────────────────────────────────────────────
    add_info_table(doc, [
        ("Document Reference", "IT-SEC-BROWSER-001"),
        ("Version",            "1.0"),
        ("Classification",     "Internal Use Only"),
        ("Policy Owner",       "IT Security"),
        ("Review Date",        "Annually or upon significant change"),
    ])

    doc.add_paragraph()

    # ── 1. Purpose ────────────────────────────────────────────────────────────
    add_heading1(doc, "1.  Purpose")
    add_body(doc,
        "This standard establishes the approved web browsers for use on all corporate "
        "devices and defines the requirements for browser configuration, security, and "
        "exceptions. Standardisation on supported browsers ensures consistent security "
        "controls, compatibility with corporate systems, and simplified management "
        "across the organisation.")

    # ── 2. Scope ──────────────────────────────────────────────────────────────
    add_heading1(doc, "2.  Scope")
    add_body(doc, "This standard applies to:")
    add_bullet(doc, "All employees, contractors, and third parties accessing corporate systems via a web browser")
    add_bullet(doc, "All corporate-owned and managed devices, including desktops, laptops, and virtual desktop infrastructure (VDI)")
    add_bullet(doc, "Any system or service delivered via a web browser interface")

    # ── 3. Approved Browsers ──────────────────────────────────────────────────
    add_heading1(doc, "3.  Approved Browsers")
    add_body(doc,
        "The following browsers are approved for use on corporate devices. All other "
        "browsers are not permitted unless a formal exception has been granted "
        "(see Section 6).")
    doc.add_paragraph()

    add_browser_table(doc, [
        ("Microsoft Edge",       "APPROVED",      "Primary approved browser. Preferred for all corporate use.",       True),
        ("Google Chrome",        "APPROVED",      "Secondary approved browser. Permitted for corporate use.",         True),
        ("Mozilla Firefox",      "NOT PERMITTED", "Not approved for standard use. Exception required.",               False),
        ("Apple Safari",         "NOT PERMITTED", "Not approved for standard use. Exception required.",               False),
        ("Opera / Brave / Other","NOT PERMITTED", "Not approved for standard use. Exception required.",               False),
    ])

    # ── 4. Rationale ──────────────────────────────────────────────────────────
    add_heading1(doc, "4.  Rationale for Standardisation")
    add_body(doc,
        "The decision to standardise on Microsoft Edge and Google Chrome is based on "
        "the following considerations.")

    add_heading2(doc, "4.1  Microsoft Account Integration & Single Sign-On")
    add_body(doc,
        "Microsoft Edge is deeply integrated with the Microsoft 365 ecosystem and "
        "supports seamless pass-through authentication using corporate Microsoft "
        "accounts. This enables:")
    add_bullet(doc, "Automatic sign-in to Microsoft 365 services (SharePoint, Teams, Outlook Web, etc.) without repeated credential prompts")
    add_bullet(doc, "Conditional Access policy enforcement through Azure Active Directory (Azure AD)")
    add_bullet(doc, "Compliance with organisational identity and access management (IAM) controls")
    add_bullet(doc, "Reduced credential exposure by minimising password entry events")

    add_heading2(doc, "4.2  Security & Patch Management")
    add_body(doc,
        "Both approved browsers receive regular, automated security updates and are "
        "supported by enterprise management tooling. Standardisation reduces the "
        "attack surface by:")
    add_bullet(doc, "Limiting browser-based vulnerabilities to a managed and patched set of applications")
    add_bullet(doc, "Enabling centrally enforced browser configuration via Group Policy (GPO) and Intune")
    add_bullet(doc, "Providing consistent support for corporate certificate authorities and TLS inspection controls")

    add_heading2(doc, "4.3  Compatibility")
    add_body(doc,
        "Corporate web applications, internal portals, and third-party SaaS platforms "
        "are tested and supported against the approved browsers. Use of non-approved "
        "browsers may result in degraded functionality or inability to access required "
        "business systems.")

    # ── 5. Configuration ──────────────────────────────────────────────────────
    add_heading1(doc, "5.  Configuration Requirements")
    add_body(doc,
        "The following baseline configuration requirements apply to all approved "
        "browsers on corporate devices:")
    add_bullet(doc, "Browsers must be managed via Group Policy or Intune device configuration profiles")
    add_bullet(doc, "Automatic updates must remain enabled and must not be disabled by users")
    add_bullet(doc, "Browser sync must be restricted to the corporate Microsoft or Google account only")
    add_bullet(doc, "Safe Browsing / SmartScreen must be enabled at all times")
    add_bullet(doc, "Password saving in the browser must be disabled; corporate password managers are to be used")
    add_bullet(doc, "Installation of browser extensions must be restricted to IT-approved extensions only")
    add_bullet(doc, "Private / InPrivate browsing is permitted but must not be used to circumvent corporate controls")

    # ── 6. Exceptions ─────────────────────────────────────────────────────────
    add_heading1(doc, "6.  Exception Process")
    add_body(doc,
        "Where a business requirement necessitates the use of a non-approved browser, "
        "a formal exception must be raised through the IT Security team prior to "
        "installation or use. Exceptions will be assessed on a case-by-case basis "
        "and must include:")
    add_bullet(doc, "Business justification for the use of the non-approved browser")
    add_bullet(doc, "Identification of the specific system or application requiring the alternative browser")
    add_bullet(doc, "Risk assessment and any proposed mitigating controls")
    add_bullet(doc, "Approval from the relevant business owner and IT Security")
    add_body(doc,
        "Approved exceptions will be time-limited and subject to periodic review. "
        "Unauthorised use of non-approved browsers may be subject to disciplinary "
        "action in accordance with the Acceptable Use Policy.")

    # ── 7. Roles & Responsibilities ───────────────────────────────────────────
    add_heading1(doc, "7.  Roles & Responsibilities")
    add_roles_table(doc, [
        ("IT Security",
         "Own and maintain this standard; review exceptions; enforce compliance"),
        ("IT Operations / Desktop",
         "Deploy and configure approved browsers; apply Group Policy and Intune settings; manage updates"),
        ("All Staff",
         "Use only approved browsers; report any issues; seek exception approval before using alternative browsers"),
    ])

    # ── 8. Compliance ─────────────────────────────────────────────────────────
    add_heading1(doc, "8.  Compliance & Enforcement")
    add_body(doc,
        "Compliance with this standard is mandatory. IT Security and IT Operations "
        "will periodically audit devices to verify adherence. Non-compliant browsers "
        "identified during audits will be subject to removal. Repeated or deliberate "
        "non-compliance may be escalated in accordance with the organisation's "
        "disciplinary procedures.")

    # ── 9. Review ─────────────────────────────────────────────────────────────
    add_heading1(doc, "9.  Document Review")
    add_body(doc,
        "This standard will be reviewed annually or following any significant change "
        "to the organisation's browser landscape, security requirements, or relevant "
        "regulatory obligations. The IT Security team is responsible for initiating "
        "and completing the review.")

    # ── Sign-off ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    sig = doc.add_paragraph()
    add_bottom_border_to_para(sig)   # repurposed as top rule via spacing
    sig.paragraph_format.space_before = Pt(20)
    add_run(sig,
            "Approved by: ________________________     Date: _______________",
            colour=GREY, size_pt=10)

    sig2 = doc.add_paragraph()
    add_run(sig2,
            "Role: ________________________________     Signature: ___________",
            colour=GREY, size_pt=10)

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == "__main__":
    build_document()
